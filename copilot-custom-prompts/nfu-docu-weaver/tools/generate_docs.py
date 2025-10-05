import os
import yaml
import re
import json
import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import argparse
from pathlib import Path
from tqdm import tqdm
from dataclasses import dataclass, field
from typing import List, Dict, Set, Any, Optional
from datetime import datetime
from enum import Enum


# ============================================================================
# æ•°æ®ç±»å®šä¹‰ (Data Classes)
# ============================================================================

class FieldType(Enum):
    """å­—æ®µç±»å‹æšä¸¾"""
    STRING = "string"
    INTEGER = "integer"
    FLOAT = "float"
    BOOLEAN = "boolean"
    LIST = "list"
    DICT = "dict"
    UNKNOWN = "unknown"


@dataclass
class ValidationRules:
    """éªŒè¯è§„åˆ™æ•°æ®ç±»
    
    ä» Schema æå–çš„éªŒè¯è§„åˆ™é›†åˆ
    
    Attributes:
        required_fields: å¿…éœ€å­—æ®µé›†åˆï¼ˆå­—æ®µè·¯å¾„ï¼‰
        optional_fields: å¯é€‰å­—æ®µé›†åˆï¼ˆå­—æ®µè·¯å¾„ï¼‰
        field_types: å­—æ®µç±»å‹æ˜ å°„ {å­—æ®µè·¯å¾„: FieldType}
        nested_structures: åµŒå¥—ç»“æ„å®šä¹‰ {å­—æ®µè·¯å¾„: å­ç»“æ„}
        list_fields: åˆ—è¡¨ç±»å‹å­—æ®µé›†åˆ
    """
    required_fields: Set[str] = field(default_factory=set)
    optional_fields: Set[str] = field(default_factory=set)
    field_types: Dict[str, FieldType] = field(default_factory=dict)
    nested_structures: Dict[str, Any] = field(default_factory=dict)
    list_fields: Set[str] = field(default_factory=set)


@dataclass
class ValidationError:
    """éªŒè¯é”™è¯¯æ•°æ®ç±»
    
    Attributes:
        field_path: å­—æ®µè·¯å¾„ï¼ˆå¦‚ "class_hours.total"ï¼‰
        error_type: é”™è¯¯ç±»å‹ï¼ˆå¦‚ "missing_required_field", "type_mismatch"ï¼‰
        message: é”™è¯¯æ¶ˆæ¯
        severity: ä¸¥é‡çº§åˆ«ï¼ˆ"error" æˆ– "warning"ï¼‰
    """
    field_path: str
    error_type: str
    message: str
    severity: str = "error"
    
    def to_dict(self) -> dict:
        """è½¬æ¢ä¸ºå­—å…¸æ ¼å¼"""
        return {
            "type": self.error_type,
            "field": self.field_path,
            "message": self.message,
            "severity": self.severity
        }


@dataclass
class ValidationResult:
    """éªŒè¯ç»“æœæ•°æ®ç±»
    
    Attributes:
        file_path: è¢«éªŒè¯æ–‡ä»¶çš„è·¯å¾„
        is_valid: æ˜¯å¦é€šè¿‡éªŒè¯ï¼ˆæ— é”™è¯¯ï¼‰
        errors: é”™è¯¯åˆ—è¡¨
        warnings: è­¦å‘Šåˆ—è¡¨
        timestamp: éªŒè¯æ—¶é—´æˆ³
        schema_version: ä½¿ç”¨çš„ Schema ç‰ˆæœ¬
        checked_fields: å·²æ£€æŸ¥çš„å­—æ®µæ•°
    """
    file_path: str
    is_valid: bool
    errors: List[ValidationError] = field(default_factory=list)
    warnings: List[ValidationError] = field(default_factory=list)
    timestamp: datetime = field(default_factory=datetime.now)
    schema_version: str = "1.0.0"
    checked_fields: int = 0
    
    def to_dict(self) -> dict:
        """è½¬æ¢ä¸ºå­—å…¸æ ¼å¼ï¼ˆç”¨äº JSON è¾“å‡ºï¼‰"""
        return {
            "file": self.file_path,
            "valid": self.is_valid,
            "errors": [e.to_dict() for e in self.errors],
            "warnings": [w.to_dict() for w in self.warnings],
            "timestamp": self.timestamp.isoformat(),
            "schema_version": self.schema_version,
            "checked_fields": self.checked_fields
        }


# ============================================================================
# å¼‚å¸¸ç±»å®šä¹‰ (Exception Classes)
# ============================================================================

class SchemaValidationError(Exception):
    """Schema éªŒè¯ç›¸å…³å¼‚å¸¸åŸºç±»"""
    pass


class SchemaLoadError(SchemaValidationError):
    """Schema åŠ è½½å¤±è´¥å¼‚å¸¸"""
    pass


class SchemaVersionError(SchemaValidationError):
    """Schema ç‰ˆæœ¬ä¸å…¼å®¹å¼‚å¸¸"""
    pass


class SchemaNotLoadedError(SchemaValidationError):
    """Schema æœªåŠ è½½å°±å°è¯•éªŒè¯å¼‚å¸¸"""
    pass


class DataParseError(SchemaValidationError):
    """æ•°æ®è§£æå¤±è´¥å¼‚å¸¸"""
    pass


# ============================================================================
# SchemaValidator ç±»
# ============================================================================

class SchemaValidator:
    """åŸºäº Schema çš„éªŒè¯å™¨ - ä» Schema æ–‡ä»¶æå–è§„åˆ™å¹¶éªŒè¯æ•°æ®
    
    æ ¸å¿ƒèŒè´£ï¼š
    - åŠ è½½å¹¶è§£æ Schema å®šä¹‰
    - ä» Schema æå–éªŒè¯è§„åˆ™
    - æ‰§è¡Œæ•°æ®éªŒè¯
    - ç”Ÿæˆè¯¦ç»†çš„éªŒè¯æŠ¥å‘Š
    """
    
    def __init__(self, schema_path: Optional[str] = None):
        """åˆå§‹åŒ– Schema éªŒè¯å™¨
        
        Args:
            schema_path: Schema æ–‡ä»¶è·¯å¾„ï¼ˆå¯é€‰ï¼Œå¯åç»­é€šè¿‡ load_schema åŠ è½½ï¼‰
        """
        self.schema_path: Optional[Path] = Path(schema_path) if schema_path else None
        self.schema: Optional[dict] = None
        self.validation_rules: Optional[ValidationRules] = None
        self._schema_cache: Dict[str, dict] = {}  # Schema ç¼“å­˜
        
        if schema_path:
            self.load_schema(schema_path)
    
    def load_schema(self, schema_path: str) -> dict:
        """åŠ è½½ Schema å®šä¹‰
        
        Args:
            schema_path: Schema YAML æ–‡ä»¶è·¯å¾„
            
        Returns:
            dict: è§£æåçš„ Schema æ•°æ®
            
        Raises:
            SchemaLoadError: Schema åŠ è½½å¤±è´¥
        """
        schema_path = Path(schema_path)
        
        # æ£€æŸ¥ç¼“å­˜
        cache_key = str(schema_path.resolve())
        if cache_key in self._schema_cache:
            self.schema = self._schema_cache[cache_key]
            self.schema_path = schema_path
            self.validation_rules = self.extract_rules(self.schema)
            return self.schema
        
        if not schema_path.exists():
            raise SchemaLoadError(f"Schema æ–‡ä»¶ä¸å­˜åœ¨: {schema_path}")
        
        try:
            with open(schema_path, 'r', encoding='utf-8') as f:
                self.schema = yaml.safe_load(f)
            
            if self.schema is None:
                raise SchemaLoadError("Schema æ–‡ä»¶ä¸ºç©º")
            
            if not isinstance(self.schema, dict):
                raise SchemaLoadError("Schema æ ¹èŠ‚ç‚¹å¿…é¡»æ˜¯å­—å…¸ç±»å‹")
            
            self.schema_path = schema_path
            # ç¼“å­˜ Schema
            self._schema_cache[cache_key] = self.schema
            # åŠ è½½åç«‹å³æå–è§„åˆ™
            self.validation_rules = self.extract_rules(self.schema)
            
            return self.schema
            
        except yaml.YAMLError as e:
            raise SchemaLoadError(f"Schema æ–‡ä»¶æ ¼å¼é”™è¯¯: {str(e)}")
    
    def extract_rules(self, schema: dict) -> ValidationRules:
        """ä» Schema æå–éªŒè¯è§„åˆ™
        
        Args:
            schema: Schema å­—å…¸
            
        Returns:
            ValidationRules: éªŒè¯è§„åˆ™å¯¹è±¡ï¼ŒåŒ…å«ï¼š
                - required_fields: å¿…éœ€å­—æ®µé›†åˆ
                - field_types: å­—æ®µç±»å‹æ˜ å°„
                - nested_structures: åµŒå¥—ç»“æ„å®šä¹‰
                - list_fields: åˆ—è¡¨ç±»å‹å­—æ®µ
        """
        rules = ValidationRules()
        
        # é€’å½’æå–æ‰€æœ‰å­—æ®µå’Œç±»å‹ä¿¡æ¯
        self._extract_fields_recursive(schema, "", rules)
        
        return rules
    
    def _extract_fields_recursive(self, data: Any, prefix: str, rules: ValidationRules) -> None:
        """é€’å½’æå–å­—æ®µä¿¡æ¯
        
        Args:
            data: å½“å‰å±‚çº§çš„æ•°æ®
            prefix: å­—æ®µè·¯å¾„å‰ç¼€ï¼ˆå¦‚ "class_hours" æˆ– "design_details"ï¼‰
            rules: ValidationRules å¯¹è±¡ï¼ˆä¼šè¢«ä¿®æ”¹ï¼‰
        """
        if not isinstance(data, dict):
            return
        
        for key, value in data.items():
            # è·³è¿‡æ³¨é‡Šå­—æ®µï¼ˆä»¥ # å¼€å¤´çš„é”®ä¸æ˜¯å®é™…æ•°æ®å­—æ®µï¼‰
            if key.startswith('#'):
                continue
            
            # æ„å»ºå®Œæ•´çš„å­—æ®µè·¯å¾„
            full_key = f"{prefix}.{key}" if prefix else key
            
            # æ‰€æœ‰åœ¨ schema ä¸­å‡ºç°çš„å­—æ®µéƒ½è§†ä¸ºå¿…éœ€å­—æ®µ
            # ï¼ˆé™¤éå€¼ä¸º None æˆ–ç©ºå­—ç¬¦ä¸²ï¼Œè¿™äº›å¯èƒ½æ˜¯å¯é€‰çš„ï¼‰
            if value is not None and value != "":
                rules.required_fields.add(full_key)
            else:
                rules.optional_fields.add(full_key)
            
            # ç¡®å®šå­—æ®µç±»å‹
            if isinstance(value, dict):
                rules.field_types[full_key] = FieldType.DICT
                rules.nested_structures[full_key] = value
                # é€’å½’å¤„ç†åµŒå¥—ç»“æ„
                self._extract_fields_recursive(value, full_key, rules)
            elif isinstance(value, list):
                rules.field_types[full_key] = FieldType.LIST
                rules.list_fields.add(full_key)
                # å¦‚æœåˆ—è¡¨åŒ…å«å­—å…¸ï¼Œå¤„ç†ç¬¬ä¸€ä¸ªå…ƒç´ ä½œä¸ºæ¨¡æ¿
                if value and isinstance(value[0], dict):
                    rules.nested_structures[full_key] = value[0]
                    # ä¸ºåˆ—è¡¨é¡¹åˆ›å»ºæ¨¡æ¿éªŒè¯è§„åˆ™
                    self._extract_fields_recursive(value[0], f"{full_key}[]", rules)
            elif isinstance(value, bool):
                # æ³¨æ„ï¼šbool æ£€æŸ¥å¿…é¡»åœ¨ int ä¹‹å‰ï¼Œå› ä¸º bool æ˜¯ int çš„å­ç±»
                rules.field_types[full_key] = FieldType.BOOLEAN
            elif isinstance(value, int):
                rules.field_types[full_key] = FieldType.INTEGER
            elif isinstance(value, float):
                rules.field_types[full_key] = FieldType.FLOAT
            elif isinstance(value, str):
                rules.field_types[full_key] = FieldType.STRING
            else:
                rules.field_types[full_key] = FieldType.UNKNOWN
    
    def validate_against_schema(self, data: dict, file_path: str = "", schema: Optional[dict] = None) -> ValidationResult:
        """æ ¹æ® Schema éªŒè¯æ•°æ®
        
        Args:
            data: è¦éªŒè¯çš„æ•°æ®ï¼ˆå­—å…¸ï¼‰
            file_path: æ•°æ®æ–‡ä»¶è·¯å¾„ï¼ˆç”¨äºç»“æœå±•ç¤ºï¼‰
            schema: Schema å®šä¹‰ï¼ˆå¯é€‰ï¼Œé»˜è®¤ä½¿ç”¨å·²åŠ è½½çš„ schemaï¼‰
            
        Returns:
            ValidationResult: éªŒè¯ç»“æœå¯¹è±¡
            
        Raises:
            SchemaNotLoadedError: Schema æœªåŠ è½½
        """
        if schema is None:
            if self.schema is None:
                raise SchemaNotLoadedError("æœªåŠ è½½ Schemaï¼Œè¯·å…ˆè°ƒç”¨ load_schema()")
            schema = self.schema
        
        # å¦‚æœä¼ å…¥äº†æ–°çš„ schemaï¼Œé‡æ–°æå–è§„åˆ™
        if schema != self.schema:
            rules = self.extract_rules(schema)
        else:
            rules = self.validation_rules
        
        # åˆ›å»ºéªŒè¯ç»“æœå¯¹è±¡
        result = ValidationResult(
            file_path=file_path,
            is_valid=True,
            errors=[],
            warnings=[]
        )
        
        # 1. æ£€æŸ¥å¿…éœ€å­—æ®µæ˜¯å¦å­˜åœ¨
        data_fields = self._get_all_data_fields(data)
        
        # è¿‡æ»¤æ‰åˆ—è¡¨æ¨¡æ¿å­—æ®µï¼ˆåŒ…å« [] çš„ï¼‰
        required_fields_filtered = {
            f for f in rules.required_fields 
            if "[]" not in f
        }
        
        missing_fields = required_fields_filtered - data_fields
        
        for field in missing_fields:
            error = ValidationError(
                field_path=field,
                error_type="missing_required_field",
                message=f"ç¼ºå°‘å¿…éœ€å­—æ®µ: '{field}'",
                severity="error"
            )
            result.errors.append(error)
            result.is_valid = False
        
        # 2. æ£€æŸ¥å­—æ®µç±»å‹æ˜¯å¦åŒ¹é…
        for field in data_fields:
            result.checked_fields += 1
            
            # è·³è¿‡ä¸åœ¨ schema ä¸­çš„å­—æ®µï¼ˆä¼šåœ¨åé¢çš„é¢å¤–å­—æ®µæ£€æŸ¥ä¸­å¤„ç†ï¼‰
            if field not in rules.field_types:
                continue
            
            expected_type = rules.field_types[field]
            actual_value = self._get_nested_value(data, field)
            
            # ç±»å‹éªŒè¯
            type_valid, type_error = self._validate_field_type(
                field, actual_value, expected_type
            )
            
            if not type_valid:
                error = ValidationError(
                    field_path=field,
                    error_type="type_mismatch",
                    message=type_error,
                    severity="error"
                )
                result.errors.append(error)
                result.is_valid = False
        
        # 3. æ£€æŸ¥é¢å¤–å­—æ®µï¼ˆå¯èƒ½çš„æ‹¼å†™é”™è¯¯ï¼‰
        extra_fields = data_fields - required_fields_filtered - rules.optional_fields
        
        for field in extra_fields:
            # æ£€æŸ¥æ˜¯å¦æ˜¯åˆ—è¡¨é¡¹çš„å­—æ®µï¼ˆå¦‚ main_teaching_segments[0].segment_titleï¼‰
            # è¿™äº›ä¸åº”è¯¥è¢«æ ‡è®°ä¸ºé¢å¤–å­—æ®µ
            if self._is_list_item_field(field, rules.list_fields):
                continue
            
            warning = ValidationError(
                field_path=field,
                error_type="extra_field",
                message=f"æ•°æ®ä¸­å­˜åœ¨ Schema æœªå®šä¹‰çš„å­—æ®µ: '{field}' (å¯èƒ½æ˜¯æ‹¼å†™é”™è¯¯æˆ–å¤šä½™å­—æ®µ)",
                severity="warning"
            )
            result.warnings.append(warning)
        
        # 4. éªŒè¯åµŒå¥—ç»“æ„å®Œæ•´æ€§ï¼ˆé’ˆå¯¹åˆ—è¡¨ç±»å‹ï¼‰
        for list_field in rules.list_fields:
            if list_field in data_fields:
                list_value = self._get_nested_value(data, list_field)
                if isinstance(list_value, list) and list_value:
                    # éªŒè¯åˆ—è¡¨ä¸­çš„æ¯ä¸ªå¯¹è±¡
                    self._validate_list_items(
                        list_field, list_value, rules, result
                    )
        
        return result
    
    def _get_all_data_fields(self, data, prefix=''):
        """é€’å½’è·å–æ•°æ®ä¸­çš„æ‰€æœ‰å­—æ®µè·¯å¾„"""
        fields = set()
        
        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{prefix}.{key}" if prefix else key
                fields.add(full_key)
                
                if isinstance(value, dict):
                    fields.update(self._get_all_data_fields(value, full_key))
                elif isinstance(value, list):
                    # å¯¹äºåˆ—è¡¨ï¼Œæ£€æŸ¥ç¬¬ä¸€ä¸ªå…ƒç´ çš„ç»“æ„
                    if value and isinstance(value[0], dict):
                        # æ·»åŠ åˆ—è¡¨é¡¹çš„å­—æ®µ
                        for item in value:
                            if isinstance(item, dict):
                                fields.update(self._get_all_data_fields(item, full_key))
        
        return fields
    
    def _get_nested_value(self, data, key_path):
        """è·å–åµŒå¥—å­—å…¸ä¸­çš„å€¼"""
        keys = key_path.split('.')
        value = data
        
        try:
            for key in keys:
                if isinstance(value, dict):
                    value = value.get(key)
                    if value is None:
                        return None
                else:
                    return None
            return value
        except (KeyError, TypeError):
            return None
    
    def _validate_field_type(self, field: str, value: Any, expected_type: FieldType) -> tuple[bool, Optional[str]]:
        """éªŒè¯å­—æ®µç±»å‹
        
        Args:
            field: å­—æ®µè·¯å¾„
            value: å®é™…å€¼
            expected_type: æœŸæœ›çš„å­—æ®µç±»å‹ï¼ˆFieldType æšä¸¾ï¼‰
            
        Returns:
            tuple: (is_valid, error_message)
        """
        if value is None:
            return True, None  # None å€¼åœ¨å¿…éœ€å­—æ®µæ£€æŸ¥ä¸­å¤„ç†
        
        # æ˜ å°„ FieldType æšä¸¾åˆ° Python ç±»å‹
        type_checks = {
            FieldType.STRING: (str, "å­—ç¬¦ä¸²"),
            FieldType.INTEGER: (int, "æ•´æ•°"),
            FieldType.FLOAT: ((int, float), "æ•°å­—"),
            FieldType.BOOLEAN: (bool, "å¸ƒå°”å€¼"),
            FieldType.LIST: (list, "åˆ—è¡¨"),
            FieldType.DICT: (dict, "å¯¹è±¡/å­—å…¸"),
        }
        
        if expected_type in type_checks:
            expected_python_type, type_name = type_checks[expected_type]
            if not isinstance(value, expected_python_type):
                actual_type = type(value).__name__
                return False, f"å­—æ®µ '{field}' ç±»å‹é”™è¯¯: æœŸæœ› {type_name}ï¼Œå®é™…ä¸º {actual_type}"
        
        return True, None
    
    def _is_list_item_field(self, field, list_fields):
        """åˆ¤æ–­å­—æ®µæ˜¯å¦æ˜¯åˆ—è¡¨é¡¹çš„å­å­—æ®µ"""
        for list_field in list_fields:
            if field.startswith(list_field + "."):
                return True
        return False
    
    def _validate_list_items(self, list_field: str, list_value: list, rules: ValidationRules, result: ValidationResult) -> None:
        """éªŒè¯åˆ—è¡¨é¡¹çš„ç»“æ„å®Œæ•´æ€§
        
        Args:
            list_field: åˆ—è¡¨å­—æ®µè·¯å¾„
            list_value: åˆ—è¡¨å€¼
            rules: éªŒè¯è§„åˆ™å¯¹è±¡
            result: éªŒè¯ç»“æœå¯¹è±¡ï¼ˆä¼šè¢«ä¿®æ”¹ï¼‰
        """
        template_key = f"{list_field}[]"
        
        # è·å–åˆ—è¡¨é¡¹åº”è¯¥æœ‰çš„å­—æ®µ
        expected_item_fields = set()
        for field in rules.required_fields:
            if field.startswith(template_key + "."):
                # æå–å­—æ®µåï¼ˆå»æ‰ list_field[] å‰ç¼€ï¼‰
                item_field = field.replace(template_key + ".", "")
                expected_item_fields.add(item_field)
        
        # å¦‚æœæ²¡æœ‰å®šä¹‰åˆ—è¡¨é¡¹ç»“æ„ï¼Œè·³è¿‡
        if not expected_item_fields:
            return
        
        # éªŒè¯æ¯ä¸ªåˆ—è¡¨é¡¹
        for idx, item in enumerate(list_value):
            if not isinstance(item, dict):
                continue
            
            item_fields = set(item.keys())
            missing_in_item = expected_item_fields - item_fields
            
            for missing_field in missing_in_item:
                warning = ValidationError(
                    field_path=f"{list_field}[{idx}].{missing_field}",
                    error_type="incomplete_list_item",
                    message=f"åˆ—è¡¨é¡¹ {list_field}[{idx}] ç¼ºå°‘å­—æ®µ: '{missing_field}'",
                    severity="warning"
                )
                result.warnings.append(warning)


class DataValidator:
    """æ•°æ®éªŒè¯å™¨ - éªŒè¯ YAML æ•°æ®æ–‡ä»¶çš„å®Œæ•´æ€§"""
    
    def __init__(self, schema_path=None):
        """åˆå§‹åŒ–æ•°æ®éªŒè¯å™¨
        
        Args:
            schema_path: Schema æ–‡ä»¶è·¯å¾„ï¼ˆå¯é€‰ï¼‰ã€‚å¦‚æœæä¾›ï¼Œå°†ä½¿ç”¨ schema éªŒè¯
        """
        self.errors = []
        self.warnings = []
        self.schema_validator = None
        
        # å¦‚æœæä¾›äº† schema_pathï¼Œåˆå§‹åŒ– SchemaValidator
        if schema_path:
            try:
                self.schema_validator = SchemaValidator(schema_path)
            except Exception as e:
                # Schema åŠ è½½å¤±è´¥ä¸å½±å“åŸºæœ¬éªŒè¯åŠŸèƒ½
                print(f"è­¦å‘Š: æ— æ³•åŠ è½½ Schema: {str(e)}")
                self.schema_validator = None
    
    def validate_yaml_syntax(self, file_path):
        """éªŒè¯ YAML æ–‡ä»¶è¯­æ³•æ˜¯å¦æ­£ç¡®
        
        Returns:
            tuple: (is_valid, data_or_error)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
                if data is None:
                    return False, "YAML æ–‡ä»¶ä¸ºç©º"
                if not isinstance(data, dict):
                    return False, "YAML æ ¹èŠ‚ç‚¹å¿…é¡»æ˜¯å­—å…¸ç±»å‹"
                return True, data
        except yaml.YAMLError as e:
            return False, f"YAML è¯­æ³•é”™è¯¯: {str(e)}"
        except FileNotFoundError:
            return False, "æ–‡ä»¶ä¸å­˜åœ¨"
        except Exception as e:
            return False, f"è¯»å–æ–‡ä»¶é”™è¯¯: {str(e)}"
    
    def extract_placeholders(self, doc):
        """ä»æ¨¡æ¿æ–‡æ¡£ä¸­æå–æ‰€æœ‰å ä½ç¬¦
        
        Args:
            doc: python-docx Document å¯¹è±¡
            
        Returns:
            set: å ä½ç¬¦é”®åé›†åˆï¼ˆä¸å« {{}} ç¬¦å·ï¼‰
        """
        placeholders = set()
        pattern = r'\{\{([^}]+)\}\}'
        
        # ä»æ®µè½ä¸­æå–
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # ä»è¡¨æ ¼ä¸­æå–
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
        
        return placeholders
    
    def validate_required_keys(self, data, required_keys):
        """éªŒè¯æ•°æ®ä¸­æ˜¯å¦åŒ…å«æ‰€æœ‰å¿…éœ€çš„é”®
        
        Args:
            data: YAML æ•°æ®ï¼ˆå­—å…¸ï¼‰
            required_keys: å¿…éœ€çš„é”®é›†åˆ
            
        Returns:
            tuple: (is_valid, missing_keys, extra_keys)
        """
        data_keys = set(self._get_all_keys(data))
        missing_keys = required_keys - data_keys
        extra_keys = data_keys - required_keys
        
        is_valid = len(missing_keys) == 0
        return is_valid, missing_keys, extra_keys
    
    def _get_all_keys(self, data, prefix=''):
        """é€’å½’è·å–æ‰€æœ‰é”®ï¼ˆåŒ…æ‹¬åµŒå¥—é”®ï¼Œä½¿ç”¨ç‚¹å·è¡¨ç¤ºï¼‰"""
        keys = []
        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{prefix}.{key}" if prefix else key
                keys.append(full_key)
                if isinstance(value, dict):
                    keys.extend(self._get_all_keys(value, full_key))
        return keys
    
    def validate_single_file(self, data_file, template_path):
        """éªŒè¯å•ä¸ªæ•°æ®æ–‡ä»¶
        
        Returns:
            dict: éªŒè¯ç»“æœ
        """
        result = {
            "file": str(data_file),
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # 1. éªŒè¯ YAML è¯­æ³•
        is_valid, data_or_error = self.validate_yaml_syntax(data_file)
        if not is_valid:
            result["valid"] = False
            result["errors"].append({
                "type": "yaml_syntax",
                "message": data_or_error,
                "severity": "error"
            })
            return result
        
        data = data_or_error
        
        # 2. å¦‚æœæœ‰ SchemaValidatorï¼Œä½¿ç”¨ schema éªŒè¯ï¼ˆä¼˜å…ˆï¼‰
        if self.schema_validator:
            try:
                # è°ƒç”¨ SchemaValidator è¿›è¡ŒéªŒè¯
                schema_result = self.schema_validator.validate_against_schema(
                    data, 
                    file_path=str(data_file)
                )
                
                # å°† ValidationResult è½¬æ¢ä¸ºå­—å…¸å¹¶åˆå¹¶ç»“æœ
                result["errors"].extend([e.to_dict() for e in schema_result.errors])
                result["warnings"].extend([w.to_dict() for w in schema_result.warnings])
                
                if not schema_result.is_valid:
                    result["valid"] = False
                
                # æ·»åŠ  schema éªŒè¯çš„é¢å¤–ä¿¡æ¯
                result["schema_validation"] = {
                    "enabled": True,
                    "checked_fields": schema_result.checked_fields,
                    "schema_version": schema_result.schema_version,
                    "timestamp": schema_result.timestamp.isoformat()
                }
                
            except SchemaNotLoadedError as e:
                result["warnings"].append({
                    "type": "schema_not_loaded",
                    "message": f"Schema æœªåŠ è½½: {str(e)}",
                    "severity": "warning"
                })
            except SchemaValidationError as e:
                result["warnings"].append({
                    "type": "schema_validation_error",
                    "message": f"Schema éªŒè¯å¤±è´¥: {str(e)}",
                    "severity": "warning"
                })
            except Exception as e:
                result["warnings"].append({
                    "type": "unexpected_error",
                    "message": f"éªŒè¯è¿‡ç¨‹å‘ç”Ÿé”™è¯¯: {str(e)}",
                    "severity": "warning"
                })
        else:
            # 3. å›é€€åˆ°åŸºäºæ¨¡æ¿çš„éªŒè¯ï¼ˆå‘åå…¼å®¹ï¼‰
            try:
                from docx import Document
                doc = Document(template_path)
                required_keys = self.extract_placeholders(doc)
            except Exception as e:
                result["valid"] = False
                result["errors"].append({
                    "type": "template_error",
                    "message": f"æ— æ³•è¯»å–æ¨¡æ¿: {str(e)}",
                    "severity": "error"
                })
                return result
            
            # éªŒè¯å¿…éœ€é”®
            is_valid, missing_keys, extra_keys = self.validate_required_keys(data, required_keys)
            
            if missing_keys:
                result["valid"] = False
                for key in missing_keys:
                    result["errors"].append({
                        "type": "missing_key",
                        "message": f"ç¼ºå°‘å¿…éœ€çš„é”®: '{key}'",
                        "key": key,
                        "severity": "error"
                    })
            
            if extra_keys:
                for key in extra_keys:
                    result["warnings"].append({
                        "type": "extra_key",
                        "message": f"æœªä½¿ç”¨çš„æ•°æ®é”®: '{key}'",
                        "key": key,
                        "severity": "warning"
                    })
            
            result["schema_validation"] = {
                "enabled": False
            }
        
        return result
    
    def validate_batch(self, data_dir, template_path):
        """æ‰¹é‡éªŒè¯æ•°æ®æ–‡ä»¶
        
        Returns:
            dict: æ‰¹é‡éªŒè¯ç»“æœ
        """
        dir_path = Path(data_dir)
        yaml_files = sorted(list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml')))
        
        if not yaml_files:
            return {
                "success": False,
                "error": f"æœªåœ¨ç›®å½•ä¸­æ‰¾åˆ° YAML æ–‡ä»¶: {data_dir}",
                "total": 0,
                "valid": 0,
                "invalid": 0
            }
        
        results = []
        valid_count = 0
        invalid_count = 0
        
        for data_file in yaml_files:
            result = self.validate_single_file(data_file, template_path)
            results.append(result)
            
            if result["valid"]:
                valid_count += 1
            else:
                invalid_count += 1
        
        return {
            "success": invalid_count == 0,
            "total": len(yaml_files),
            "valid": valid_count,
            "invalid": invalid_count,
            "results": results
        }


class DocumentGenerator:
    def __init__(self):
        self.debug = False
    
    def log(self, message):
        """è°ƒè¯•æ—¥å¿—è¾“å‡º"""
        if self.debug:
            print(f"[DEBUG] {message}")
    
    def copy_paragraph_format(self, source_paragraph, target_paragraph):
        """å®‰å…¨åœ°å¤åˆ¶æ®µè½æ ¼å¼"""
        try:
            # å¤åˆ¶æ®µè½æ ¼å¼
            if source_paragraph.paragraph_format:
                target_format = target_paragraph.paragraph_format
                source_format = source_paragraph.paragraph_format
                
                # å¤åˆ¶å¯¹é½æ–¹å¼
                if source_format.alignment is not None:
                    target_format.alignment = source_format.alignment
                
                # å¤åˆ¶ç¼©è¿›
                if source_format.left_indent is not None:
                    target_format.left_indent = source_format.left_indent
                if source_format.right_indent is not None:
                    target_format.right_indent = source_format.right_indent
                if source_format.first_line_indent is not None:
                    target_format.first_line_indent = source_format.first_line_indent
                
                # å¤åˆ¶é—´è·
                if source_format.space_before is not None:
                    target_format.space_before = source_format.space_before
                if source_format.space_after is not None:
                    target_format.space_after = source_format.space_after
                if source_format.line_spacing is not None:
                    target_format.line_spacing = source_format.line_spacing
            
            # å¤åˆ¶æ ·å¼
            if source_paragraph.style:
                target_paragraph.style = source_paragraph.style
                
            self.log(f"æˆåŠŸå¤åˆ¶æ®µè½æ ¼å¼")
            
        except Exception as e:
            self.log(f"å¤åˆ¶æ®µè½æ ¼å¼æ—¶å‡ºé”™: {str(e)}")
    
    def get_base_run_format(self, paragraph):
        """è·å–æ®µè½çš„åŸºç¡€æ–‡å­—æ ¼å¼"""
        base_format = {}
        
        try:
            # ä¼˜å…ˆä»ç°æœ‰runè·å–æ ¼å¼ï¼ˆæ›´å‡†ç¡®ï¼‰
            if paragraph.runs:
                first_run = paragraph.runs[0]
                
                # å­—ä½“åç§°å’Œå¤§å°
                if first_run.font.name:
                    base_format['font_name'] = first_run.font.name
                if first_run.font.size:
                    base_format['font_size'] = first_run.font.size
                
                # è·å–é¢œè‰²ï¼ˆæ›´å¥å£®çš„æ–¹å¼ï¼‰
                if first_run.font.color:
                    if first_run.font.color.rgb:
                        base_format['font_color'] = first_run.font.color.rgb
                    elif first_run.font.color.theme_color:
                        # å¦‚æœæ˜¯ä¸»é¢˜é¢œè‰²ï¼Œä¹Ÿå°è¯•è·å–
                        base_format['font_color_theme'] = first_run.font.color.theme_color
                
                # ä¿ç•™ç²—ä½“ã€æ–œä½“ã€ä¸‹åˆ’çº¿çŠ¶æ€ï¼ˆåŒ…æ‹¬ Noneï¼‰
                base_format['font_bold'] = first_run.font.bold
                base_format['font_italic'] = first_run.font.italic
                base_format['font_underline'] = first_run.font.underline
                
                # ä¿ç•™å…¶ä»–æ ¼å¼
                base_format['font_strike'] = first_run.font.strike
                base_format['font_all_caps'] = first_run.font.all_caps
                base_format['font_small_caps'] = first_run.font.small_caps
            
            # å¦‚æœä»runè·å–ä¸åˆ°ï¼Œå†ä»æ®µè½æ ·å¼è·å–
            if not base_format and paragraph.style and hasattr(paragraph.style, 'font'):
                style_font = paragraph.style.font
                if style_font.name:
                    base_format['font_name'] = style_font.name
                if style_font.size:
                    base_format['font_size'] = style_font.size
                if style_font.color and style_font.color.rgb:
                    base_format['font_color'] = style_font.color.rgb
            
            self.log(f"è·å–åˆ°çš„åŸºç¡€æ ¼å¼: {base_format}")
        
        except Exception as e:
            self.log(f"è·å–åŸºç¡€æ ¼å¼æ—¶å‡ºé”™: {str(e)}")
        
        return base_format
    
    def apply_run_format(self, run, base_format):
        """åº”ç”¨åŸºç¡€æ ¼å¼åˆ°run"""
        try:
            # å­—ä½“åç§°å’Œå¤§å°
            if 'font_name' in base_format:
                font_name = base_format['font_name']
                run.font.name = font_name
                
                # ğŸ”‘ å…³é”®ä¿®å¤ï¼šåŒæ—¶è®¾ç½®ä¸œäºšå­—ä½“ï¼ˆä¸­æ–‡ï¼‰å’Œå¤æ‚è„šæœ¬å­—ä½“
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                
                # è®¾ç½®æ‰€æœ‰å­—ä½“ç±»å‹
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)  # ä¸œäºšå­—ä½“ï¼ˆä¸­æ–‡ï¼‰
                rFonts.set(qn('w:cs'), font_name)        # å¤æ‚è„šæœ¬å­—ä½“
            
            if 'font_size' in base_format:
                run.font.size = base_format['font_size']
            
            # é¢œè‰²
            if 'font_color' in base_format:
                run.font.color.rgb = base_format['font_color']
            
            # ğŸ”‘ å…³é”®ä¿®å¤ï¼šä¸ç»§æ‰¿ç²—ä½“ã€æ–œä½“ã€ä¸‹åˆ’çº¿æ ¼å¼ï¼Œé¿å…æ ¼å¼æ±¡æŸ“
            # æ³¨é‡Šæ‰ä»¥ä¸‹ä»£ç ï¼Œé˜²æ­¢æ›¿æ¢æ–‡æœ¬ç»§æ‰¿å‰é¢çš„æ ¼å¼
            # if 'font_bold' in base_format:
            #     run.font.bold = base_format['font_bold']
            # if 'font_italic' in base_format:
            #     run.font.italic = base_format['font_italic']
            # if 'font_underline' in base_format:
            #     run.font.underline = base_format['font_underline']
            
            # å…¶ä»–æ ¼å¼
            if 'font_strike' in base_format:
                run.font.strike = base_format['font_strike']
            if 'font_all_caps' in base_format:
                run.font.all_caps = base_format['font_all_caps']
            if 'font_small_caps' in base_format:
                run.font.small_caps = base_format['font_small_caps']
                
        except Exception as e:
            self.log(f"åº”ç”¨runæ ¼å¼æ—¶å‡ºé”™: {str(e)}")
    
    def parse_markdown_formatting(self, text):
        """è§£æMarkdownæ ¼å¼å¹¶è¿”å›æ ¼å¼åŒ–çš„ç‰‡æ®µåˆ—è¡¨"""
        segments = []
        
        # å¤„ç†ç²—ä½“å’Œæ–œä½“çš„æ­£åˆ™è¡¨è¾¾å¼
        # æ”¯æŒ **ç²—ä½“** å’Œ *æ–œä½“*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # ç²—ä½“æ–‡æœ¬
                segments.append({
                    'text': part[2:-2],  # å»æ‰ ** æ ‡è®°
                    'bold': True,
                    'italic': False
                })
            elif part.startswith('*') and part.endswith('*'):
                # æ–œä½“æ–‡æœ¬
                segments.append({
                    'text': part[1:-1],  # å»æ‰ * æ ‡è®°
                    'bold': False,
                    'italic': True
                })
            else:
                # æ™®é€šæ–‡æœ¬
                segments.append({
                    'text': part,
                    'bold': False,
                    'italic': False
                })
        
        return segments
    
    def add_formatted_text_to_paragraph(self, paragraph, text, base_format):
        """å‘æ®µè½æ·»åŠ æ ¼å¼åŒ–æ–‡æœ¬"""
        # ğŸ”‘ å…³é”®ä¿®å¤ï¼šåœ¨æ¸…ç©ºå†…å®¹å‰ä¿å­˜åŸå§‹æ®µè½æ ¼å¼
        original_format = None
        if paragraph.paragraph_format:
            original_format = {
                'left_indent': paragraph.paragraph_format.left_indent,
                'right_indent': paragraph.paragraph_format.right_indent,
                'first_line_indent': paragraph.paragraph_format.first_line_indent,
                'alignment': paragraph.paragraph_format.alignment,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing
            }
        
        # æ¸…ç©ºç°æœ‰å†…å®¹
        paragraph.clear()
        
        # ğŸ”‘ å…³é”®ä¿®å¤ï¼šæ¢å¤æ®µè½æ ¼å¼ï¼Œç¡®ä¿ç¼©è¿›ä¿æŒä¸€è‡´
        if original_format and paragraph.paragraph_format:
            target_format = paragraph.paragraph_format
            if original_format['left_indent'] is not None:
                target_format.left_indent = original_format['left_indent']
            if original_format['right_indent'] is not None:
                target_format.right_indent = original_format['right_indent']
            if original_format['first_line_indent'] is not None:
                target_format.first_line_indent = original_format['first_line_indent']
            if original_format['alignment'] is not None:
                target_format.alignment = original_format['alignment']
            if original_format['space_before'] is not None:
                target_format.space_before = original_format['space_before']
            if original_format['space_after'] is not None:
                target_format.space_after = original_format['space_after']
            if original_format['line_spacing'] is not None:
                target_format.line_spacing = original_format['line_spacing']
        
        # ğŸ”‘ å…³é”®ä¿®å¤ï¼šå¤„ç†æ¢è¡Œç¬¦çš„æ­£ç¡®æ–¹å¼
        # å¦‚æœæ–‡æœ¬åŒ…å«æ¢è¡Œç¬¦ï¼Œéœ€è¦ç‰¹æ®Šå¤„ç†ä»¥ä¿æŒç¼©è¿›
        # æ³¨æ„ï¼šè¿™é‡Œæ£€æµ‹çš„æ˜¯å®é™…çš„æ¢è¡Œç¬¦ï¼Œä¸æ˜¯å­—ç¬¦ä¸² '\n'
        if '\n' in text:
            # ğŸ”‘ æ–°æ€è·¯ï¼šå°†å¤šè¡Œæ–‡æœ¬åˆ†å‰²æˆå¤šä¸ªç‹¬ç«‹çš„æ®µè½
            # æ¯ä¸ªæ®µè½éƒ½æœ‰è‡ªå·±çš„é¦–è¡Œç¼©è¿›ï¼Œè¿™æ ·èƒ½ç¡®ä¿æ¯è¡Œéƒ½æ­£ç¡®ç¼©è¿›
            lines = text.split('\n')
            
            # å¤„ç†ç¬¬ä¸€è¡Œ
            first_line = lines[0]
            segments = self.parse_markdown_formatting(first_line)
            
            for segment in segments:
                run = paragraph.add_run(segment['text'])
                
                # åº”ç”¨åŸºç¡€æ ¼å¼
                self.apply_run_format(run, base_format)
                
                # åº”ç”¨Markdownæ ¼å¼ï¼ˆåªåœ¨æœ‰ Markdown æ ‡è®°æ—¶è¦†ç›–ï¼‰
                if segment['bold']:
                    run.font.bold = True
                if segment['italic']:
                    run.font.italic = True
            
            # å¤„ç†åç»­è¡Œï¼Œæ¯è¡Œåˆ›å»ºä¸€ä¸ªæ–°æ®µè½
            for line in lines[1:]:
                if line.strip():  # è·³è¿‡ç©ºè¡Œ
                    # ğŸ”‘ å…³é”®ä¿®å¤ï¼šåœ¨è¡¨æ ¼å•å…ƒæ ¼ä¸­æ­£ç¡®åˆ›å»ºæ–°æ®µè½
                    # è·å–å½“å‰æ®µè½çš„çˆ¶å®¹å™¨
                    parent = paragraph._element.getparent()
                    
                    # åˆ›å»ºæ–°æ®µè½å…ƒç´ 
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    new_para_elem = OxmlElement('w:p')
                    parent.append(new_para_elem)
                    
                    # åˆ›å»ºæ®µè½å¯¹è±¡
                    from docx.text.paragraph import Paragraph
                    new_para = Paragraph(new_para_elem, parent)
                    
                    # å¤åˆ¶æ®µè½æ ¼å¼
                    if original_format and new_para.paragraph_format:
                        target_format = new_para.paragraph_format
                        if original_format['left_indent'] is not None:
                            target_format.left_indent = original_format['left_indent']
                        if original_format['right_indent'] is not None:
                            target_format.right_indent = original_format['right_indent']
                        if original_format['first_line_indent'] is not None:
                            target_format.first_line_indent = original_format['first_line_indent']
                        if original_format['alignment'] is not None:
                            target_format.alignment = original_format['alignment']
                        if original_format['space_before'] is not None:
                            target_format.space_before = original_format['space_before']
                        if original_format['space_after'] is not None:
                            target_format.space_after = original_format['space_after']
                        if original_format['line_spacing'] is not None:
                            target_format.line_spacing = original_format['line_spacing']
                    
                    # æ·»åŠ æ–‡æœ¬åˆ°æ–°æ®µè½
                    segments = self.parse_markdown_formatting(line)
                    
                    for segment in segments:
                        run = new_para.add_run(segment['text'])
                        
                        # åº”ç”¨åŸºç¡€æ ¼å¼
                        self.apply_run_format(run, base_format)
                        
                        # åº”ç”¨Markdownæ ¼å¼ï¼ˆåªåœ¨æœ‰ Markdown æ ‡è®°æ—¶è¦†ç›–ï¼‰
                        if segment['bold']:
                            run.font.bold = True
                        if segment['italic']:
                            run.font.italic = True
        else:
            # æ²¡æœ‰æ¢è¡Œç¬¦ï¼Œæ­£å¸¸å¤„ç†
            segments = self.parse_markdown_formatting(text)
            
            for segment in segments:
                run = paragraph.add_run(segment['text'])
                
                # åº”ç”¨åŸºç¡€æ ¼å¼
                self.apply_run_format(run, base_format)
                
                # åº”ç”¨Markdownæ ¼å¼ï¼ˆåªåœ¨æœ‰ Markdown æ ‡è®°æ—¶è¦†ç›–ï¼‰
                if segment['bold']:
                    run.font.bold = True
                if segment['italic']:
                    run.font.italic = True
    
    def process_template_paragraph(self, paragraph, data):
        """å¤„ç†å•ä¸ªæ¨¡æ¿æ®µè½"""
        if not paragraph.text.strip():
            return False
        
        original_text = paragraph.text
        self.log(f"å¤„ç†æ®µè½: {original_text[:50]}...")
        
        # ğŸ”‘ å…³é”®ä¿®å¤ï¼šæ£€æŸ¥æ˜¯å¦åŒ…å«å ä½ç¬¦
        if '{{' not in original_text or '}}' not in original_text:
            return False
        
        # ğŸ”‘ å…³é”®ä¿®å¤ï¼šåªæ›¿æ¢å ä½ç¬¦ï¼Œä¿ç•™å‰åæ–‡å­—å’Œæ ¼å¼
        return self.replace_placeholders_in_runs(paragraph, data)
    
    def replace_placeholders_in_runs(self, paragraph, data):
        """åœ¨æ®µè½ä¸­æ›¿æ¢å ä½ç¬¦ï¼Œç²¾ç¡®ä¿ç•™æ¯ä¸ªrunçš„æ ¼å¼
        
        å…³é”®æ”¹è¿›ï¼š
        1. å¤„ç†è·¨runçš„å ä½ç¬¦ï¼ˆWordç»å¸¸å°†{{key}}æ‹†åˆ†æˆå¤šä¸ªrunï¼‰
        2. ä¿ç•™æ¯ä¸ªrunçš„ç‹¬ç«‹æ ¼å¼ï¼ˆç²—ä½“ã€æ–œä½“ç­‰ï¼‰
        3. å ä½ç¬¦å‰åçš„æ–‡å­—æ ¼å¼ä¸å—å½±å“
        4. å ä½ç¬¦æ›¿æ¢å€¼ä½¿ç”¨å…¶è‡ªèº«æ ¼å¼ï¼Œä¸ç»§æ‰¿å‰é¢çš„è£…é¥°æ€§æ ¼å¼
        """
        import re
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        
        placeholder_pattern = r'\{\{([^}]+)\}\}'
        
        # ä¿å­˜æ®µè½æ ¼å¼
        original_para_format = None
        if paragraph.paragraph_format:
            original_para_format = {
                'left_indent': paragraph.paragraph_format.left_indent,
                'right_indent': paragraph.paragraph_format.right_indent,
                'first_line_indent': paragraph.paragraph_format.first_line_indent,
                'alignment': paragraph.paragraph_format.alignment,
                'space_before': paragraph.paragraph_format.space_before,
                'space_after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing
            }
        
        # å…ˆåˆå¹¶æ‰€æœ‰runçš„æ–‡æœ¬ï¼Œæ£€æŸ¥æ˜¯å¦åŒ…å«å ä½ç¬¦
        full_text = paragraph.text
        if not re.search(placeholder_pattern, full_text):
            return False
        
        self.log(f"å¤„ç†æ®µè½: {full_text[:100]}...")
        
        # æ”¶é›†æ‰€æœ‰runåŠå…¶æ ¼å¼ä¿¡æ¯å’Œæ–‡æœ¬ä½ç½®
        runs_info = []
        current_pos = 0
        for run in paragraph.runs:
            run_text = run.text
            run_format = {
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_bold': run.font.bold,
                'font_italic': run.font.italic,
                'font_underline': run.font.underline,
                'font_color': run.font.color.rgb if run.font.color.rgb else None,
            }
            runs_info.append({
                'text': run_text,
                'format': run_format,
                'start_pos': current_pos,
                'end_pos': current_pos + len(run_text)
            })
            current_pos += len(run_text)
        
        # æ‰¾åˆ°æ‰€æœ‰å ä½ç¬¦çš„ä½ç½®
        placeholder_matches = list(re.finditer(placeholder_pattern, full_text))
        
        # æ„å»ºæ–°çš„runæ•°æ®
        new_runs_data = []
        last_processed_pos = 0
        
        for match in placeholder_matches:
            match_start = match.start()
            match_end = match.end()
            key_path = match.group(1)
            
            # 1. æ·»åŠ å ä½ç¬¦å‰çš„æ–‡æœ¬ï¼ˆä¿ç•™åŸæ ¼å¼ï¼‰
            if match_start > last_processed_pos:
                before_text = full_text[last_processed_pos:match_start]
                # æ‰¾åˆ°è¿™æ®µæ–‡æœ¬å¯¹åº”çš„runæ ¼å¼
                for run_info in runs_info:
                    if run_info['start_pos'] <= last_processed_pos < run_info['end_pos']:
                        # ä»è¿™ä¸ªrunå¼€å§‹
                        segment_start = last_processed_pos
                        while segment_start < match_start:
                            # æ‰¾åˆ°å½“å‰ä½ç½®å¯¹åº”çš„run
                            for ri in runs_info:
                                if ri['start_pos'] <= segment_start < ri['end_pos']:
                                    # æå–è¿™ä¸ªrunèŒƒå›´å†…çš„æ–‡æœ¬
                                    segment_end = min(match_start, ri['end_pos'])
                                    segment_text = full_text[segment_start:segment_end]
                                    if segment_text:
                                        new_runs_data.append({
                                            'text': segment_text,
                                            'format': ri['format'].copy()
                                        })
                                    segment_start = segment_end
                                    break
                        break
            
            # 2. æ·»åŠ æ›¿æ¢å€¼ï¼ˆä½¿ç”¨å ä½ç¬¦ä½ç½®çš„æ ¼å¼ï¼Œä½†å»é™¤è£…é¥°ï¼‰
            replacement_value = self.get_nested_value(data, key_path)
            
            if replacement_value:
                # æ‰¾åˆ°å ä½ç¬¦æ‰€åœ¨ä½ç½®çš„æ ¼å¼ï¼ˆå–å ä½ç¬¦å¼€å§‹ä½ç½®çš„æ ¼å¼ï¼‰
                placeholder_format = None
                for run_info in runs_info:
                    if run_info['start_pos'] <= match_start < run_info['end_pos']:
                        placeholder_format = run_info['format'].copy()
                        break
                
                # å¦‚æœæ²¡æ‰¾åˆ°ï¼Œä½¿ç”¨æœ€åä¸€ä¸ªéç©ºrunçš„æ ¼å¼
                if not placeholder_format:
                    for run_info in reversed(runs_info):
                        if run_info['format'].get('font_name'):
                            placeholder_format = run_info['format'].copy()
                            break
                
                # ç¡®ä¿æœ‰æ ¼å¼
                if not placeholder_format:
                    placeholder_format = {}
                
                # å»é™¤è£…é¥°æ€§æ ¼å¼
                placeholder_format['font_bold'] = False
                placeholder_format['font_italic'] = False
                placeholder_format['font_underline'] = False
                
                # å¤„ç†æ¢è¡Œç¬¦
                if '\n' in replacement_value:
                    lines = replacement_value.split('\n')
                    
                    # ç¬¬ä¸€è¡Œæ·»åŠ åˆ°å½“å‰æ®µè½
                    if lines[0]:
                        new_runs_data.append({
                            'text': lines[0],
                            'format': placeholder_format,
                            'is_replacement': True
                        })
                    
                    # åç»­è¡Œåˆ›å»ºæ–°æ®µè½
                    for line in lines[1:]:
                        if line.strip():
                            new_runs_data.append({
                                'text': line,
                                'format': placeholder_format,
                                'is_replacement': True,
                                'new_paragraph': True
                            })
                else:
                    new_runs_data.append({
                        'text': replacement_value,
                        'format': placeholder_format,
                        'is_replacement': True
                    })
            
            last_processed_pos = match_end
        
        # 3. æ·»åŠ æœ€åä¸€ä¸ªå ä½ç¬¦ä¹‹åçš„æ–‡æœ¬
        if last_processed_pos < len(full_text):
            after_text = full_text[last_processed_pos:]
            segment_start = last_processed_pos
            while segment_start < len(full_text):
                for ri in runs_info:
                    if ri['start_pos'] <= segment_start < ri['end_pos']:
                        segment_end = min(len(full_text), ri['end_pos'])
                        segment_text = full_text[segment_start:segment_end]
                        if segment_text:
                            new_runs_data.append({
                                'text': segment_text,
                                'format': ri['format'].copy()
                            })
                        segment_start = segment_end
                        break
        
        # é‡å»ºæ®µè½å†…å®¹
        paragraph.clear()
        
        # æ¢å¤æ®µè½æ ¼å¼
        if original_para_format and paragraph.paragraph_format:
            target_format = paragraph.paragraph_format
            for key, value in original_para_format.items():
                if value is not None:
                    setattr(target_format, key, value)
        
        # æ·»åŠ runåˆ°æ®µè½
        current_paragraph = paragraph
        
        for run_data in new_runs_data:
            # å¦‚æœéœ€è¦åˆ›å»ºæ–°æ®µè½
            if run_data.get('new_paragraph'):
                parent = paragraph._element.getparent()
                new_para_elem = OxmlElement('w:p')
                parent.append(new_para_elem)
                current_paragraph = Paragraph(new_para_elem, parent)
                
                # å¤åˆ¶æ®µè½æ ¼å¼
                if original_para_format and current_paragraph.paragraph_format:
                    target_format = current_paragraph.paragraph_format
                    for key, value in original_para_format.items():
                        if value is not None:
                            setattr(target_format, key, value)
            
            # æ·»åŠ run
            run = current_paragraph.add_run(run_data['text'])
            
            # åº”ç”¨æ ¼å¼
            run_format = run_data['format']
            if run_format.get('font_name'):
                run.font.name = run_format['font_name']
                
                # è®¾ç½®æ‰€æœ‰å­—ä½“ç±»å‹ï¼ˆç¡®ä¿ä¸­æ–‡æ˜¾ç¤ºæ­£ç¡®ï¼‰
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                
                font_name = run_format['font_name']
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:cs'), font_name)
            
            if run_format.get('font_size'):
                run.font.size = run_format['font_size']
            
            if run_format.get('font_color'):
                run.font.color.rgb = run_format['font_color']
            
            # åªåœ¨æ˜ç¡®ä¸ºTrueæ—¶è®¾ç½®ç²—ä½“ã€æ–œä½“ã€ä¸‹åˆ’çº¿
            if run_format.get('font_bold') is True:
                run.font.bold = True
            elif run_format.get('font_bold') is False:
                run.font.bold = False
            
            if run_format.get('font_italic') is True:
                run.font.italic = True
            elif run_format.get('font_italic') is False:
                run.font.italic = False
            
            if run_format.get('font_underline'):
                run.font.underline = run_format['font_underline']
        
        self.log(f"æ›¿æ¢å®Œæˆï¼Œå…±åˆ›å»º {len([r for r in new_runs_data if not r.get('new_paragraph')])} ä¸ªrun")
        return True
    
    def replace_text_in_runs(self, paragraph, start_pos, end_pos, replacement_value, base_format):
        """åœ¨æŒ‡å®šçš„ä½ç½®æ›¿æ¢æ–‡æœ¬"""
        # ğŸ”‘ ç®€åŒ–é€»è¾‘ï¼šç›´æ¥æ›¿æ¢å ä½ç¬¦æ–‡æœ¬
        # æ‰¾åˆ°åŒ…å«å ä½ç¬¦çš„ run å¹¶æ›¿æ¢
        current_pos = 0
        
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if run_start <= start_pos and run_end >= end_pos:
                # è¿™ä¸ª run åŒ…å«å ä½ç¬¦
                before_text = run.text[:start_pos - run_start]
                after_text = run.text[end_pos - run_start:]
                
                # æ›¿æ¢æ–‡æœ¬
                run.text = before_text + (replacement_value or '') + after_text
                break
            
            current_pos = run_end
    
    def copy_run_format(self, source_run, target_run):
        """å¤åˆ¶ run çš„æ ¼å¼"""
        try:
            # å¤åˆ¶å­—ä½“æ ¼å¼
            if source_run.font.name:
                target_run.font.name = source_run.font.name
            if source_run.font.size:
                target_run.font.size = source_run.font.size
            if source_run.font.bold is not None:
                target_run.font.bold = source_run.font.bold
            if source_run.font.italic is not None:
                target_run.font.italic = source_run.font.italic
            if source_run.font.underline is not None:
                target_run.font.underline = source_run.font.underline
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
        except Exception as e:
            self.log(f"å¤åˆ¶runæ ¼å¼æ—¶å‡ºé”™: {str(e)}")
    
    def create_replacement_runs(self, paragraph, replacement_value, base_format):
        """åˆ›å»ºæ›¿æ¢æ–‡æœ¬çš„ runs"""
        runs = []
        
        if not replacement_value:
            return runs
        
        # å¦‚æœåŒ…å«æ¢è¡Œç¬¦ï¼Œéœ€è¦ç‰¹æ®Šå¤„ç†
        if '\n' in replacement_value:
            lines = replacement_value.split('\n')
            
            for i, line in enumerate(lines):
                if line.strip():  # è·³è¿‡ç©ºè¡Œ
                    # åˆ›å»ºæ–°æ®µè½ï¼ˆå¦‚æœéœ€è¦ï¼‰
                    if i > 0:
                        # è·å–å½“å‰æ®µè½çš„æ ¼å¼
                        original_format = None
                        if paragraph.paragraph_format:
                            original_format = {
                                'left_indent': paragraph.paragraph_format.left_indent,
                                'right_indent': paragraph.paragraph_format.right_indent,
                                'first_line_indent': paragraph.paragraph_format.first_line_indent,
                                'alignment': paragraph.paragraph_format.alignment,
                                'space_before': paragraph.paragraph_format.space_before,
                                'space_after': paragraph.paragraph_format.space_after,
                                'line_spacing': paragraph.paragraph_format.line_spacing
                            }
                        
                        # åˆ›å»ºæ–°æ®µè½
                        parent = paragraph._element.getparent()
                        from docx.oxml import OxmlElement
                        new_para_elem = OxmlElement('w:p')
                        parent.append(new_para_elem)
                        
                        from docx.text.paragraph import Paragraph
                        new_para = Paragraph(new_para_elem, parent)
                        
                        # å¤åˆ¶æ®µè½æ ¼å¼
                        if original_format and new_para.paragraph_format:
                            target_format = new_para.paragraph_format
                            if original_format['left_indent'] is not None:
                                target_format.left_indent = original_format['left_indent']
                            if original_format['right_indent'] is not None:
                                target_format.right_indent = original_format['right_indent']
                            if original_format['first_line_indent'] is not None:
                                target_format.first_line_indent = original_format['first_line_indent']
                            if original_format['alignment'] is not None:
                                target_format.alignment = original_format['alignment']
                            if original_format['space_before'] is not None:
                                target_format.space_before = original_format['space_before']
                            if original_format['space_after'] is not None:
                                target_format.space_after = original_format['space_after']
                            if original_format['line_spacing'] is not None:
                                target_format.line_spacing = original_format['line_spacing']
                        
                        # åœ¨æ–°æ®µè½ä¸­æ·»åŠ æ–‡æœ¬
                        run = new_para.add_run(line)
                        self.apply_run_format(run, base_format)
                        runs.append(run)
                    else:
                        # ç¬¬ä¸€è¡Œï¼Œåœ¨å½“å‰æ®µè½ä¸­æ·»åŠ 
                        run = paragraph.add_run(line)
                        self.apply_run_format(run, base_format)
                        runs.append(run)
        else:
            # æ²¡æœ‰æ¢è¡Œç¬¦ï¼Œæ­£å¸¸å¤„ç†
            run = paragraph.add_run(replacement_value)
            self.apply_run_format(run, base_format)
            runs.append(run)
        
        return runs
    
    def split_into_multiple_paragraphs(self, original_paragraph, text, base_format, data):
        """å°†æ–‡æœ¬åˆ†å‰²æˆå¤šä¸ªæ®µè½"""
        # æŒ‰ \n\n åˆ†å‰²æ®µè½
        paragraphs_text = text.split('\n\n')
        
        # å¤„ç†ç¬¬ä¸€ä¸ªæ®µè½ï¼ˆæ›¿æ¢åŸæ®µè½ï¼‰
        if paragraphs_text:
            first_text = paragraphs_text[0].strip()
            if first_text:
                self.add_formatted_text_to_paragraph(original_paragraph, first_text, base_format)
        
        # æ£€æŸ¥æ®µè½æ˜¯å¦åœ¨è¡¨æ ¼å•å…ƒæ ¼ä¸­
        parent = original_paragraph._element.getparent()
        
        # å¦‚æœæ®µè½åœ¨è¡¨æ ¼å•å…ƒæ ¼ä¸­ï¼ˆparent æ˜¯ tc å…ƒç´ ï¼‰
        if parent.tag.endswith('}tc'):
            # åœ¨è¡¨æ ¼å•å…ƒæ ¼ä¸­ï¼Œç›´æ¥åœ¨å½“å‰å•å…ƒæ ¼æ·»åŠ æ–°æ®µè½
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph
            
            for para_text in paragraphs_text[1:]:
                para_text = para_text.strip()
                if para_text:
                    # åœ¨å•å…ƒæ ¼ä¸­æ·»åŠ æ–°æ®µè½
                    new_p_element = OxmlElement('w:p')
                    parent.append(new_p_element)
                    
                    # åˆ›å»º Paragraph å¯¹è±¡
                    new_paragraph = Paragraph(new_p_element, parent)
                    
                    # å¤åˆ¶æ ¼å¼
                    self.copy_paragraph_format(original_paragraph, new_paragraph)
                    
                    # æ·»åŠ æ ¼å¼åŒ–æ–‡æœ¬
                    self.add_formatted_text_to_paragraph(new_paragraph, para_text, base_format)
        else:
            # åœ¨æ™®é€šæ–‡æ¡£æµä¸­ï¼ŒæŒ‰åŸæœ‰é€»è¾‘å¤„ç†
            original_index = list(parent).index(original_paragraph._element)
            
            for i, para_text in enumerate(paragraphs_text[1:], 1):
                para_text = para_text.strip()
                if para_text:
                    # è·å–æ–‡æ¡£å¯¹è±¡
                    if hasattr(original_paragraph, '_part'):
                        doc_obj = original_paragraph._part.document
                    else:
                        # å¦‚æœæ— æ³•è·å–æ–‡æ¡£å¯¹è±¡ï¼Œè·³è¿‡
                        self.log("æ— æ³•è·å–æ–‡æ¡£å¯¹è±¡ï¼Œè·³è¿‡æ®µè½åˆ›å»º")
                        continue
                    
                    # åˆ›å»ºæ–°æ®µè½
                    new_paragraph = doc_obj.add_paragraph()
                    
                    # å¤åˆ¶æ ¼å¼
                    self.copy_paragraph_format(original_paragraph, new_paragraph)
                    
                    # æ·»åŠ æ ¼å¼åŒ–æ–‡æœ¬
                    self.add_formatted_text_to_paragraph(new_paragraph, para_text, base_format)
                    
                    # å°†æ–°æ®µè½æ’å…¥åˆ°æ­£ç¡®ä½ç½®
                    parent.insert(original_index + i, new_paragraph._element)
        
        return True
    
    def replace_placeholders(self, text, data):
        """æ›¿æ¢æ–‡æœ¬ä¸­çš„å ä½ç¬¦"""
        def replace_match(match):
            placeholder = match.group(1)
            return self.get_nested_value(data, placeholder)
        
        # åŒ¹é… {{key}} æ ¼å¼çš„å ä½ç¬¦
        result = re.sub(r'\{\{([^}]+)\}\}', replace_match, text)
        return result
    
    def get_nested_value(self, data, key_path):
        """è·å–åµŒå¥—å­—å…¸ä¸­çš„å€¼ï¼Œæ”¯æŒæ•°ç»„ç´¢å¼•"""
        # å¤„ç†æ•°ç»„ç´¢å¼•ï¼Œå¦‚ main_teaching_segments[0].design_details.organization
        keys = []
        parts = key_path.strip().split('.')
        
        for part in parts:
            if '[' in part and ']' in part:
                # å¤„ç†æ•°ç»„ç´¢å¼•
                field_name = part.split('[')[0]
                index_str = part.split('[')[1].split(']')[0]
                try:
                    index = int(index_str)
                    keys.append((field_name, index))
                except ValueError:
                    return f'{{{{{key_path}}}}}'
            else:
                keys.append(part)
        
        value = data
        
        try:
            for key in keys:
                if isinstance(key, tuple):
                    # å¤„ç†æ•°ç»„ç´¢å¼•
                    field_name, index = key
                    if isinstance(value, dict) and field_name in value:
                        if isinstance(value[field_name], list) and 0 <= index < len(value[field_name]):
                            value = value[field_name][index]
                        else:
                            return f'{{{{{key_path}}}}}'
                    else:
                        return f'{{{{{key_path}}}}}'
                else:
                    # å¤„ç†æ™®é€šå­—æ®µ
                    if isinstance(value, dict):
                        value = value.get(key, f'{{{{{key_path}}}}}')
                    else:
                        return f'{{{{{key_path}}}}}'
            
            # å¦‚æœå€¼æ˜¯åˆ—è¡¨ï¼Œè½¬æ¢ä¸ºæ¢è¡Œåˆ†éš”çš„å­—ç¬¦ä¸²
            if isinstance(value, list):
                return '\n'.join(str(item) for item in value)
            
            # å¤„ç†å­—ç¬¦ä¸²å€¼
            if isinstance(value, str):
                # å¤„ç†è½¬ä¹‰ç¬¦ï¼šå°† \n è½¬æ¢ä¸ºå®é™…æ¢è¡Œç¬¦
                processed_value = value.replace('\\n', '\n')
                return processed_value
            
            # å¤„ç†å…¶ä»–ç±»å‹
            if value is not None:
                # å¦‚æœæ˜¯å­—å…¸ï¼Œå°è¯•æå– value å­—æ®µ
                if isinstance(value, dict) and 'value' in value:
                    return str(value['value']).replace('\\n', '\n')
                return str(value).replace('\\n', '\n')
            
            return f'{{{{{key_path}}}}}'
        
        except Exception as e:
            self.log(f"è·å–å€¼æ—¶å‡ºé”™ {key_path}: {str(e)}")
            return f'{{{{{key_path}}}}}'
    
    def analyze_directory(self, directory_path, recursive=False):
        """åˆ†æç›®å½•ä¸­çš„YAMLæ–‡ä»¶æ•°é‡"""
        try:
            dir_path = Path(directory_path)
            
            if not dir_path.exists():
                return {
                    "success": False,
                    "error": f"ç›®å½•ä¸å­˜åœ¨: {directory_path}"
                }
            
            if not dir_path.is_dir():
                return {
                    "success": False,
                    "error": f"è·¯å¾„ä¸æ˜¯ç›®å½•: {directory_path}"
                }
            
            # æŸ¥æ‰¾YAMLæ–‡ä»¶
            if recursive:
                yaml_files = list(dir_path.rglob('*.yml')) + list(dir_path.rglob('*.yaml'))
            else:
                yaml_files = list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml'))
            
            return {
                "success": True,
                "file_count": len(yaml_files),
                "files": [str(f.relative_to(dir_path)) for f in yaml_files]
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def batch_generate(self, data_dir, template_path, output_dir, continue_on_error=False, debug=False, quiet=False, verbose=False):
        """æ‰¹é‡ç”Ÿæˆæ–‡æ¡£
        
        Args:
            data_dir: æ•°æ®æ–‡ä»¶ç›®å½•
            template_path: æ¨¡æ¿æ–‡ä»¶è·¯å¾„
            output_dir: è¾“å‡ºç›®å½•
            continue_on_error: é‡é”™ç»§ç»­
            debug: è°ƒè¯•æ¨¡å¼
            quiet: å®‰é™æ¨¡å¼ï¼ˆä»…æ˜¾ç¤ºé”™è¯¯ï¼‰
            verbose: è¯¦ç»†æ¨¡å¼ï¼ˆæ˜¾ç¤ºè¯¦ç»†æ—¥å¿—ï¼‰
        """
        self.debug = debug
        
        # ç¡®ä¿è¾“å‡ºç›®å½•å­˜åœ¨
        os.makedirs(output_dir, exist_ok=True)
        
        # è·å–æ‰€æœ‰YAMLæ–‡ä»¶
        dir_path = Path(data_dir)
        yaml_files = sorted(list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml')))
        
        if not yaml_files:
            return {
                "success": False,
                "error": f"æœªåœ¨ç›®å½•ä¸­æ‰¾åˆ°YAMLæ–‡ä»¶: {data_dir}",
                "total": 0,
                "succeeded": 0,
                "failed": 0
            }
        
        # æ‰¹é‡å¤„ç†
        results = []
        succeeded = 0
        failed = 0
        
        # è¿›åº¦æ¡è®¾ç½®
        use_progress_bar = not quiet and not verbose
        
        if not quiet:
            print(f"å¼€å§‹æ‰¹é‡ç”Ÿæˆæ–‡æ¡£...")
            print(f"æ€»å…±å‘ç° {len(yaml_files)} ä¸ªæ•°æ®æ–‡ä»¶")
            if not use_progress_bar:
                print("-" * 60)
        
        # ä½¿ç”¨ tqdm è¿›åº¦æ¡æˆ–æ™®é€šå¾ªç¯
        if use_progress_bar:
            pbar = tqdm(yaml_files, desc="æ‰¹é‡ç”Ÿæˆæ–‡æ¡£", unit="file")
            iterator = pbar
        else:
            iterator = yaml_files
            pbar = None
        
        for idx, data_file in enumerate(iterator, 1):
            file_name = data_file.stem
            output_path = os.path.join(output_dir, f"{file_name}.docx")
            
            if verbose and not use_progress_bar:
                print(f"\n[{idx}/{len(yaml_files)}] å¤„ç†: {data_file.name}")
            elif not use_progress_bar and not quiet:
                print(f"\n[{idx}/{len(yaml_files)}] å¤„ç†: {data_file.name}")
            
            try:
                success = self.generate_document(
                    template_path,
                    str(data_file),
                    output_path,
                    debug=False  # æ‰¹é‡æ¨¡å¼ä¸‹å…³é—­è¯¦ç»†æ—¥å¿—
                )
                
                if success:
                    succeeded += 1
                    results.append({
                        "file": data_file.name,
                        "status": "success",
                        "output": output_path
                    })
                    # Only show success message in verbose mode without progress bar
                    if verbose and not use_progress_bar:
                        print(f"  âœ… æˆåŠŸ: {output_path}")
                else:
                    failed += 1
                    results.append({
                        "file": data_file.name,
                        "status": "failed",
                        "error": "ç”Ÿæˆå¤±è´¥"
                    })
                    if not quiet:
                        print(f"  âŒ å¤±è´¥: {data_file.name}")
                    if not continue_on_error:
                        break
                        
            except Exception as e:
                failed += 1
                error_msg = str(e)
                if not quiet:
                    print(f"  âŒ é”™è¯¯: {error_msg}")
                results.append({
                    "file": data_file.name,
                    "status": "failed",
                    "error": error_msg
                })
                
                if not continue_on_error:
                    break
        
        # ç”Ÿæˆæ±‡æ€»æŠ¥å‘Š
        if not quiet:
            print("\n" + "=" * 60)
            print("æ‰¹é‡ç”Ÿæˆå®Œæˆ!")
            print(f"  âœ… æˆåŠŸ: {succeeded}")
            print(f"  âŒ å¤±è´¥: {failed}")
            print(f"  ğŸ“Š æ€»è®¡: {len(yaml_files)}")
            print("=" * 60)
        
        return {
            "success": failed == 0,
            "total": len(yaml_files),
            "succeeded": succeeded,
            "failed": failed,
            "results": results
        }
    
    def generate_document(self, template_path, data_path, output_path, debug=False):
        """ç”Ÿæˆæ–‡æ¡£çš„ä¸»å‡½æ•°"""
        self.debug = debug
        
        try:
            # åŠ è½½æ•°æ®
            self.log(f"åŠ è½½æ•°æ®æ–‡ä»¶: {data_path}")
            with open(data_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            # åŠ è½½æ¨¡æ¿
            self.log(f"åŠ è½½æ¨¡æ¿æ–‡ä»¶: {template_path}")
            doc = Document(template_path)
            
            # å¤„ç†æ‰€æœ‰æ®µè½
            processed_count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                if self.process_template_paragraph(paragraph, data):
                    processed_count += 1
            
            self.log(f"å…±å¤„ç†äº† {processed_count} ä¸ªæ®µè½")
            
            # å¤„ç†æ‰€æœ‰è¡¨æ ¼
            table_cell_count = 0
            for table_idx, table in enumerate(doc.tables):
                self.log(f"å¤„ç†è¡¨æ ¼ {table_idx}")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        # å¤„ç†å•å…ƒæ ¼ä¸­çš„æ¯ä¸ªæ®µè½
                        for paragraph in cell.paragraphs:
                            if self.process_template_paragraph(paragraph, data):
                                table_cell_count += 1
            
            self.log(f"å…±å¤„ç†äº† {table_cell_count} ä¸ªè¡¨æ ¼å•å…ƒæ ¼")
            
            # ä¿å­˜æ–‡æ¡£
            self.log(f"ä¿å­˜æ–‡æ¡£åˆ°: {output_path}")
            doc.save(output_path)
            
            if debug:
                print(f"æ–‡æ¡£å·²æˆåŠŸç”Ÿæˆ: {output_path}")
                print(f"  - å¤„ç†æ®µè½: {processed_count}")
                print(f"  - å¤„ç†è¡¨æ ¼å•å…ƒæ ¼: {table_cell_count}")
            
            return True
            
        except Exception as e:
            if debug:
                print(f"ç”Ÿæˆæ–‡æ¡£æ—¶å‡ºé”™: {str(e)}")
            return False

def cmd_validate(args):
    """validate å‘½ä»¤å¤„ç†å‡½æ•°"""
    # ç¡®å®š schema è·¯å¾„
    schema_path = getattr(args, 'schema', None)
    
    # å¦‚æœæ²¡æœ‰æŒ‡å®š schemaï¼Œå°è¯•ä½¿ç”¨é»˜è®¤è·¯å¾„
    if not schema_path:
        default_schema = Path(__file__).parent / "schemas" / "lesson_data_schema.yml"
        if default_schema.exists():
            schema_path = str(default_schema)
    
    # åˆå§‹åŒ–éªŒè¯å™¨ï¼ˆå¸¦æˆ–ä¸å¸¦ schemaï¼‰
    validator = DataValidator(schema_path=schema_path)
    
    # è¾“å‡ºéªŒè¯æ¨¡å¼ä¿¡æ¯
    if schema_path and validator.schema_validator:
        print(f"â„¹ï¸  ä½¿ç”¨ Schema éªŒè¯: {schema_path}")
    else:
        print("â„¹ï¸  ä½¿ç”¨æ¨¡æ¿å ä½ç¬¦éªŒè¯ï¼ˆæœªå¯ç”¨ Schema éªŒè¯ï¼‰")
    
    if args.data_file:
        # éªŒè¯å•ä¸ªæ–‡ä»¶
        result = validator.validate_single_file(args.data_file, args.template)
        
        # è¾“å‡ºç»“æœ
        if result["valid"]:
            print(f"âœ… éªŒè¯é€šè¿‡: {args.data_file}")
            if result["warnings"]:
                print(f"\nâš ï¸  è­¦å‘Š ({len(result['warnings'])}):")
                for warning in result["warnings"]:
                    print(f"  - {warning['message']}")
        else:
            print(f"âŒ éªŒè¯å¤±è´¥: {args.data_file}")
            print(f"\né”™è¯¯ ({len(result['errors'])}):")
            for error in result["errors"]:
                print(f"  - {error['message']}")
            
            if result["warnings"]:
                print(f"\nâš ï¸  è­¦å‘Š ({len(result['warnings'])}):")
                for warning in result["warnings"]:
                    print(f"  - {warning['message']}")
        
        # JSON è¾“å‡º
        print("\n" + json.dumps(result, ensure_ascii=False, indent=2))
        
        return 0 if result["valid"] else 1
    
    else:
        # æ‰¹é‡éªŒè¯
        result = validator.validate_batch(args.data_dir, args.template)
        
        if not result.get("success", False) and "error" in result:
            print(f"âŒ é”™è¯¯: {result['error']}")
            print(json.dumps(result, ensure_ascii=False, indent=2))
            return 1
        
        # è¾“å‡ºæ±‡æ€»
        print("=" * 60)
        print("æ•°æ®éªŒè¯å®Œæˆ!")
        print(f"  âœ… æœ‰æ•ˆ: {result['valid']}")
        print(f"  âŒ æ— æ•ˆ: {result['invalid']}")
        print(f"  ğŸ“Š æ€»è®¡: {result['total']}")
        print("=" * 60)
        
        # æ˜¾ç¤ºè¯¦ç»†é”™è¯¯
        if result['invalid'] > 0:
            print("\nè¯¦ç»†é”™è¯¯:")
            for file_result in result['results']:
                if not file_result['valid']:
                    print(f"\nâŒ {file_result['file']}:")
                    for error in file_result['errors']:
                        print(f"   - {error['message']}")
        
        # JSON è¾“å‡º
        print("\n" + "=" * 60)
        print("è¯¦ç»†ç»“æœ (JSON):")
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
        return 0 if result['success'] else 1

def cmd_analyze(args):
    """analyze å‘½ä»¤å¤„ç†å‡½æ•°"""
    generator = DocumentGenerator()
    result = generator.analyze_directory(args.directory, recursive=args.recursive)
    
    # è¾“å‡ºJSONæ ¼å¼ç»“æœ
    print(json.dumps(result, ensure_ascii=False, indent=2))
    
    return 0 if result['success'] else 1

def cmd_generate(args):
    """generate å‘½ä»¤å¤„ç†å‡½æ•°"""
    generator = DocumentGenerator()
    
    # ç¡®å®šè¾“å‡ºè·¯å¾„
    if args.output:
        output_path = args.output
    else:
        template_path = Path(args.template)
        data_path = Path(args.data)
        output_dir = args.output_dir if args.output_dir else './output'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{data_path.stem}.docx")
    
    # ç”Ÿæˆæ–‡æ¡£
    success = generator.generate_document(
        args.template, 
        args.data, 
        output_path, 
        debug=args.debug
    )
    
    # è¾“å‡ºç»“æœï¼ˆédebugæ¨¡å¼ï¼‰
    if success and not args.debug:
        print(f"æ–‡æ¡£å·²æˆåŠŸç”Ÿæˆ: {output_path}")
    elif not success:
        print(f"æ–‡æ¡£ç”Ÿæˆå¤±è´¥")
    
    # è¾“å‡ºJSONæ ¼å¼ç»“æœ
    result = {
        "success": success,
        "output_path": output_path if success else None
    }
    print(json.dumps(result, ensure_ascii=False))
    
    return 0 if success else 1

def cmd_batch(args):
    """batch å‘½ä»¤å¤„ç†å‡½æ•°"""
    generator = DocumentGenerator()
    
    # æ‰¹é‡ç”Ÿæˆ
    result = generator.batch_generate(
        args.data_dir,
        args.template,
        args.output_dir,
        continue_on_error=args.continue_on_error,
        debug=args.debug,
        quiet=args.quiet,
        verbose=args.verbose
    )
    
    # è¾“å‡ºJSONæ ¼å¼çš„è¯¦ç»†ç»“æœ
    if not args.quiet:
        print("\n" + "=" * 60)
        print("è¯¦ç»†ç»“æœ (JSON):")
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    return 0 if result['success'] else 1

def main():
    """ä¸»å‡½æ•°"""
    parser = argparse.ArgumentParser(
        description='Docu-Weaver: è‡ªåŠ¨åŒ–æ–‡æ¡£ç”Ÿæˆå·¥å…·',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
ç¤ºä¾‹:
  # åˆ†æç›®å½•ä¸­çš„YAMLæ–‡ä»¶
  python generate_docs.py analyze ./test_data
  
  # éªŒè¯å•ä¸ªæ•°æ®æ–‡ä»¶
  python generate_docs.py validate lesson1.yml template.docx
  
  # æ‰¹é‡éªŒè¯æ•°æ®æ–‡ä»¶
  python generate_docs.py validate --batch ./test_data template.docx
  
  # ç”Ÿæˆå•ä¸ªæ–‡æ¡£
  python generate_docs.py generate lesson1.yml template.docx output/
  
  # æ‰¹é‡ç”Ÿæˆæ–‡æ¡£
  python generate_docs.py batch ./test_data template.docx ./output
  
  # æ‰¹é‡ç”Ÿæˆï¼ˆé‡é”™ç»§ç»­ï¼‰
  python generate_docs.py batch ./data template.docx ./output --continue-on-error
  
  # é€’å½’åˆ†æå­ç›®å½•
  python generate_docs.py analyze ./data --recursive
        """
    )
    
    # åˆ›å»ºå­å‘½ä»¤
    subparsers = parser.add_subparsers(dest='command', help='å¯ç”¨å‘½ä»¤')
    
    # analyze å­å‘½ä»¤
    parser_analyze = subparsers.add_parser('analyze', help='åˆ†æç›®å½•ä¸­çš„YAMLæ–‡ä»¶æ•°é‡')
    parser_analyze.add_argument('directory', help='è¦åˆ†æçš„ç›®å½•è·¯å¾„')
    parser_analyze.add_argument('-r', '--recursive', action='store_true', 
                               help='é€’å½’æ‰«æå­ç›®å½•')
    parser_analyze.set_defaults(func=cmd_analyze)
    
    # validate å­å‘½ä»¤
    parser_validate = subparsers.add_parser('validate', help='éªŒè¯æ•°æ®æ–‡ä»¶å®Œæ•´æ€§')
    parser_validate.add_argument('data_file', nargs='?', help='YAMLæ•°æ®æ–‡ä»¶è·¯å¾„ï¼ˆå•æ–‡ä»¶æ¨¡å¼ï¼‰')
    parser_validate.add_argument('template', help='Wordæ¨¡æ¿æ–‡ä»¶è·¯å¾„')
    parser_validate.add_argument('--batch', dest='data_dir', help='æ•°æ®æ–‡ä»¶ç›®å½•è·¯å¾„ï¼ˆæ‰¹é‡æ¨¡å¼ï¼‰')
    parser_validate.add_argument('--schema', help='Schema æ–‡ä»¶è·¯å¾„ï¼ˆå¯é€‰ï¼Œé»˜è®¤ä½¿ç”¨ schemas/lesson_data_schema.ymlï¼‰')
    parser_validate.set_defaults(func=cmd_validate)
    
    # generate å­å‘½ä»¤
    parser_generate = subparsers.add_parser('generate', help='ç”Ÿæˆå•ä¸ªæ–‡æ¡£')
    parser_generate.add_argument('data', help='YAMLæ•°æ®æ–‡ä»¶è·¯å¾„')
    parser_generate.add_argument('template', help='Wordæ¨¡æ¿æ–‡ä»¶è·¯å¾„')
    parser_generate.add_argument('output_dir', help='è¾“å‡ºç›®å½•è·¯å¾„')
    parser_generate.add_argument('-o', '--output', help='æŒ‡å®šè¾“å‡ºæ–‡ä»¶å®Œæ•´è·¯å¾„ï¼ˆå¯é€‰ï¼‰')
    parser_generate.add_argument('--debug', action='store_true', help='å¯ç”¨è°ƒè¯•æ¨¡å¼')
    parser_generate.set_defaults(func=cmd_generate)
    
    # batch å­å‘½ä»¤
    parser_batch = subparsers.add_parser('batch', help='æ‰¹é‡ç”Ÿæˆæ–‡æ¡£')
    parser_batch.add_argument('data_dir', help='æ•°æ®æ–‡ä»¶ç›®å½•è·¯å¾„')
    parser_batch.add_argument('template', help='Wordæ¨¡æ¿æ–‡ä»¶è·¯å¾„')
    parser_batch.add_argument('output_dir', help='è¾“å‡ºç›®å½•è·¯å¾„')
    parser_batch.add_argument('--continue-on-error', action='store_true',
                             help='é‡åˆ°é”™è¯¯æ—¶ç»§ç»­å¤„ç†å…¶ä»–æ–‡ä»¶')
    parser_batch.add_argument('--debug', action='store_true', help='å¯ç”¨è°ƒè¯•æ¨¡å¼')
    parser_batch.add_argument('-q', '--quiet', action='store_true', 
                             help='å®‰é™æ¨¡å¼ï¼ˆä»…æ˜¾ç¤ºé”™è¯¯ï¼‰')
    parser_batch.add_argument('-v', '--verbose', action='store_true', 
                             help='è¯¦ç»†æ¨¡å¼ï¼ˆæ˜¾ç¤ºæ‰€æœ‰å¤„ç†ä¿¡æ¯ï¼‰')
    parser_batch.set_defaults(func=cmd_batch)
    
    # è§£æå‚æ•°
    args = parser.parse_args()
    
    # å¦‚æœæ²¡æœ‰æŒ‡å®šå­å‘½ä»¤ï¼Œæ˜¾ç¤ºå¸®åŠ©ä¿¡æ¯
    if not args.command:
        parser.print_help()
        return 1
    
    # æ‰§è¡Œå¯¹åº”çš„å‘½ä»¤å‡½æ•°
    return args.func(args)

if __name__ == "__main__":
    exit(main())
