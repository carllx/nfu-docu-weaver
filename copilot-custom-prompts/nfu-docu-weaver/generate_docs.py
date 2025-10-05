import os
import yaml
import re
import json
import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import argparse
from pathlib import Path
from tqdm import tqdm
from dataclasses import dataclass, field
from typing import List, Dict, Set, Any, Optional
from datetime import datetime
from enum import Enum


# ============================================================================
# 数据类定义 (Data Classes)
# ============================================================================

class FieldType(Enum):
    """字段类型枚举"""
    STRING = "string"
    INTEGER = "integer"
    FLOAT = "float"
    BOOLEAN = "boolean"
    LIST = "list"
    DICT = "dict"
    UNKNOWN = "unknown"


@dataclass
class ValidationRules:
    """验证规则数据类
    
    从 Schema 提取的验证规则集合
    
    Attributes:
        required_fields: 必需字段集合（字段路径）
        optional_fields: 可选字段集合（字段路径）
        field_types: 字段类型映射 {字段路径: FieldType}
        nested_structures: 嵌套结构定义 {字段路径: 子结构}
        list_fields: 列表类型字段集合
    """
    required_fields: Set[str] = field(default_factory=set)
    optional_fields: Set[str] = field(default_factory=set)
    field_types: Dict[str, FieldType] = field(default_factory=dict)
    nested_structures: Dict[str, Any] = field(default_factory=dict)
    list_fields: Set[str] = field(default_factory=set)


@dataclass
class ValidationError:
    """验证错误数据类
    
    Attributes:
        field_path: 字段路径（如 "class_hours.total"）
        error_type: 错误类型（如 "missing_required_field", "type_mismatch"）
        message: 错误消息
        severity: 严重级别（"error" 或 "warning"）
    """
    field_path: str
    error_type: str
    message: str
    severity: str = "error"
    
    def to_dict(self) -> dict:
        """转换为字典格式"""
        return {
            "type": self.error_type,
            "field": self.field_path,
            "message": self.message,
            "severity": self.severity
        }


@dataclass
class ValidationResult:
    """验证结果数据类
    
    Attributes:
        file_path: 被验证文件的路径
        is_valid: 是否通过验证（无错误）
        errors: 错误列表
        warnings: 警告列表
        timestamp: 验证时间戳
        schema_version: 使用的 Schema 版本
        checked_fields: 已检查的字段数
    """
    file_path: str
    is_valid: bool
    errors: List[ValidationError] = field(default_factory=list)
    warnings: List[ValidationError] = field(default_factory=list)
    timestamp: datetime = field(default_factory=datetime.now)
    schema_version: str = "1.0.0"
    checked_fields: int = 0
    
    def to_dict(self) -> dict:
        """转换为字典格式（用于 JSON 输出）"""
        return {
            "file": self.file_path,
            "valid": self.is_valid,
            "errors": [e.to_dict() for e in self.errors],
            "warnings": [w.to_dict() for w in self.warnings],
            "timestamp": self.timestamp.isoformat(),
            "schema_version": self.schema_version,
            "checked_fields": self.checked_fields
        }


# ============================================================================
# 异常类定义 (Exception Classes)
# ============================================================================

class SchemaValidationError(Exception):
    """Schema 验证相关异常基类"""
    pass


class SchemaLoadError(SchemaValidationError):
    """Schema 加载失败异常"""
    pass


class SchemaVersionError(SchemaValidationError):
    """Schema 版本不兼容异常"""
    pass


class SchemaNotLoadedError(SchemaValidationError):
    """Schema 未加载就尝试验证异常"""
    pass


class DataParseError(SchemaValidationError):
    """数据解析失败异常"""
    pass


# ============================================================================
# SchemaValidator 类
# ============================================================================

class SchemaValidator:
    """基于 Schema 的验证器 - 从 Schema 文件提取规则并验证数据
    
    核心职责：
    - 加载并解析 Schema 定义
    - 从 Schema 提取验证规则
    - 执行数据验证
    - 生成详细的验证报告
    """
    
    def __init__(self, schema_path: Optional[str] = None):
        """初始化 Schema 验证器
        
        Args:
            schema_path: Schema 文件路径（可选，可后续通过 load_schema 加载）
        """
        self.schema_path: Optional[Path] = Path(schema_path) if schema_path else None
        self.schema: Optional[dict] = None
        self.validation_rules: Optional[ValidationRules] = None
        self._schema_cache: Dict[str, dict] = {}  # Schema 缓存
        
        if schema_path:
            self.load_schema(schema_path)
    
    def load_schema(self, schema_path: str) -> dict:
        """加载 Schema 定义
        
        Args:
            schema_path: Schema YAML 文件路径
            
        Returns:
            dict: 解析后的 Schema 数据
            
        Raises:
            SchemaLoadError: Schema 加载失败
        """
        schema_path = Path(schema_path)
        
        # 检查缓存
        cache_key = str(schema_path.resolve())
        if cache_key in self._schema_cache:
            self.schema = self._schema_cache[cache_key]
            self.schema_path = schema_path
            self.validation_rules = self.extract_rules(self.schema)
            return self.schema
        
        if not schema_path.exists():
            raise SchemaLoadError(f"Schema 文件不存在: {schema_path}")
        
        try:
            with open(schema_path, 'r', encoding='utf-8') as f:
                self.schema = yaml.safe_load(f)
            
            if self.schema is None:
                raise SchemaLoadError("Schema 文件为空")
            
            if not isinstance(self.schema, dict):
                raise SchemaLoadError("Schema 根节点必须是字典类型")
            
            self.schema_path = schema_path
            # 缓存 Schema
            self._schema_cache[cache_key] = self.schema
            # 加载后立即提取规则
            self.validation_rules = self.extract_rules(self.schema)
            
            return self.schema
            
        except yaml.YAMLError as e:
            raise SchemaLoadError(f"Schema 文件格式错误: {str(e)}")
    
    def extract_rules(self, schema: dict) -> ValidationRules:
        """从 Schema 提取验证规则
        
        Args:
            schema: Schema 字典
            
        Returns:
            ValidationRules: 验证规则对象，包含：
                - required_fields: 必需字段集合
                - field_types: 字段类型映射
                - nested_structures: 嵌套结构定义
                - list_fields: 列表类型字段
        """
        rules = ValidationRules()
        
        # 递归提取所有字段和类型信息
        self._extract_fields_recursive(schema, "", rules)
        
        return rules
    
    def _extract_fields_recursive(self, data: Any, prefix: str, rules: ValidationRules) -> None:
        """递归提取字段信息
        
        Args:
            data: 当前层级的数据
            prefix: 字段路径前缀（如 "class_hours" 或 "design_details"）
            rules: ValidationRules 对象（会被修改）
        """
        if not isinstance(data, dict):
            return
        
        for key, value in data.items():
            # 跳过注释字段（以 # 开头的键不是实际数据字段）
            if key.startswith('#'):
                continue
            
            # 构建完整的字段路径
            full_key = f"{prefix}.{key}" if prefix else key
            
            # 所有在 schema 中出现的字段都视为必需字段
            # （除非值为 None 或空字符串，这些可能是可选的）
            if value is not None and value != "":
                rules.required_fields.add(full_key)
            else:
                rules.optional_fields.add(full_key)
            
            # 确定字段类型
            if isinstance(value, dict):
                rules.field_types[full_key] = FieldType.DICT
                rules.nested_structures[full_key] = value
                # 递归处理嵌套结构
                self._extract_fields_recursive(value, full_key, rules)
            elif isinstance(value, list):
                rules.field_types[full_key] = FieldType.LIST
                rules.list_fields.add(full_key)
                # 如果列表包含字典，处理第一个元素作为模板
                if value and isinstance(value[0], dict):
                    rules.nested_structures[full_key] = value[0]
                    # 为列表项创建模板验证规则
                    self._extract_fields_recursive(value[0], f"{full_key}[]", rules)
            elif isinstance(value, bool):
                # 注意：bool 检查必须在 int 之前，因为 bool 是 int 的子类
                rules.field_types[full_key] = FieldType.BOOLEAN
            elif isinstance(value, int):
                rules.field_types[full_key] = FieldType.INTEGER
            elif isinstance(value, float):
                rules.field_types[full_key] = FieldType.FLOAT
            elif isinstance(value, str):
                rules.field_types[full_key] = FieldType.STRING
            else:
                rules.field_types[full_key] = FieldType.UNKNOWN
    
    def validate_against_schema(self, data: dict, file_path: str = "", schema: Optional[dict] = None) -> ValidationResult:
        """根据 Schema 验证数据
        
        Args:
            data: 要验证的数据（字典）
            file_path: 数据文件路径（用于结果展示）
            schema: Schema 定义（可选，默认使用已加载的 schema）
            
        Returns:
            ValidationResult: 验证结果对象
            
        Raises:
            SchemaNotLoadedError: Schema 未加载
        """
        if schema is None:
            if self.schema is None:
                raise SchemaNotLoadedError("未加载 Schema，请先调用 load_schema()")
            schema = self.schema
        
        # 如果传入了新的 schema，重新提取规则
        if schema != self.schema:
            rules = self.extract_rules(schema)
        else:
            rules = self.validation_rules
        
        # 创建验证结果对象
        result = ValidationResult(
            file_path=file_path,
            is_valid=True,
            errors=[],
            warnings=[]
        )
        
        # 1. 检查必需字段是否存在
        data_fields = self._get_all_data_fields(data)
        
        # 过滤掉列表模板字段（包含 [] 的）
        required_fields_filtered = {
            f for f in rules.required_fields 
            if "[]" not in f
        }
        
        missing_fields = required_fields_filtered - data_fields
        
        for field in missing_fields:
            error = ValidationError(
                field_path=field,
                error_type="missing_required_field",
                message=f"缺少必需字段: '{field}'",
                severity="error"
            )
            result.errors.append(error)
            result.is_valid = False
        
        # 2. 检查字段类型是否匹配
        for field in data_fields:
            result.checked_fields += 1
            
            # 跳过不在 schema 中的字段（会在后面的额外字段检查中处理）
            if field not in rules.field_types:
                continue
            
            expected_type = rules.field_types[field]
            actual_value = self._get_nested_value(data, field)
            
            # 类型验证
            type_valid, type_error = self._validate_field_type(
                field, actual_value, expected_type
            )
            
            if not type_valid:
                error = ValidationError(
                    field_path=field,
                    error_type="type_mismatch",
                    message=type_error,
                    severity="error"
                )
                result.errors.append(error)
                result.is_valid = False
        
        # 3. 检查额外字段（可能的拼写错误）
        extra_fields = data_fields - required_fields_filtered - rules.optional_fields
        
        for field in extra_fields:
            # 检查是否是列表项的字段（如 main_teaching_segments[0].segment_title）
            # 这些不应该被标记为额外字段
            if self._is_list_item_field(field, rules.list_fields):
                continue
            
            warning = ValidationError(
                field_path=field,
                error_type="extra_field",
                message=f"数据中存在 Schema 未定义的字段: '{field}' (可能是拼写错误或多余字段)",
                severity="warning"
            )
            result.warnings.append(warning)
        
        # 4. 验证嵌套结构完整性（针对列表类型）
        for list_field in rules.list_fields:
            if list_field in data_fields:
                list_value = self._get_nested_value(data, list_field)
                if isinstance(list_value, list) and list_value:
                    # 验证列表中的每个对象
                    self._validate_list_items(
                        list_field, list_value, rules, result
                    )
        
        return result
    
    def _get_all_data_fields(self, data, prefix=''):
        """递归获取数据中的所有字段路径"""
        fields = set()
        
        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{prefix}.{key}" if prefix else key
                fields.add(full_key)
                
                if isinstance(value, dict):
                    fields.update(self._get_all_data_fields(value, full_key))
                elif isinstance(value, list):
                    # 对于列表，检查第一个元素的结构
                    if value and isinstance(value[0], dict):
                        # 添加列表项的字段
                        for item in value:
                            if isinstance(item, dict):
                                fields.update(self._get_all_data_fields(item, full_key))
        
        return fields
    
    def _get_nested_value(self, data, key_path):
        """获取嵌套字典中的值"""
        keys = key_path.split('.')
        value = data
        
        try:
            for key in keys:
                if isinstance(value, dict):
                    value = value.get(key)
                    if value is None:
                        return None
                else:
                    return None
            return value
        except (KeyError, TypeError):
            return None
    
    def _validate_field_type(self, field: str, value: Any, expected_type: FieldType) -> tuple[bool, Optional[str]]:
        """验证字段类型
        
        Args:
            field: 字段路径
            value: 实际值
            expected_type: 期望的字段类型（FieldType 枚举）
            
        Returns:
            tuple: (is_valid, error_message)
        """
        if value is None:
            return True, None  # None 值在必需字段检查中处理
        
        # 映射 FieldType 枚举到 Python 类型
        type_checks = {
            FieldType.STRING: (str, "字符串"),
            FieldType.INTEGER: (int, "整数"),
            FieldType.FLOAT: ((int, float), "数字"),
            FieldType.BOOLEAN: (bool, "布尔值"),
            FieldType.LIST: (list, "列表"),
            FieldType.DICT: (dict, "对象/字典"),
        }
        
        if expected_type in type_checks:
            expected_python_type, type_name = type_checks[expected_type]
            if not isinstance(value, expected_python_type):
                actual_type = type(value).__name__
                return False, f"字段 '{field}' 类型错误: 期望 {type_name}，实际为 {actual_type}"
        
        return True, None
    
    def _is_list_item_field(self, field, list_fields):
        """判断字段是否是列表项的子字段"""
        for list_field in list_fields:
            if field.startswith(list_field + "."):
                return True
        return False
    
    def _validate_list_items(self, list_field: str, list_value: list, rules: ValidationRules, result: ValidationResult) -> None:
        """验证列表项的结构完整性
        
        Args:
            list_field: 列表字段路径
            list_value: 列表值
            rules: 验证规则对象
            result: 验证结果对象（会被修改）
        """
        template_key = f"{list_field}[]"
        
        # 获取列表项应该有的字段
        expected_item_fields = set()
        for field in rules.required_fields:
            if field.startswith(template_key + "."):
                # 提取字段名（去掉 list_field[] 前缀）
                item_field = field.replace(template_key + ".", "")
                expected_item_fields.add(item_field)
        
        # 如果没有定义列表项结构，跳过
        if not expected_item_fields:
            return
        
        # 验证每个列表项
        for idx, item in enumerate(list_value):
            if not isinstance(item, dict):
                continue
            
            item_fields = set(item.keys())
            missing_in_item = expected_item_fields - item_fields
            
            for missing_field in missing_in_item:
                warning = ValidationError(
                    field_path=f"{list_field}[{idx}].{missing_field}",
                    error_type="incomplete_list_item",
                    message=f"列表项 {list_field}[{idx}] 缺少字段: '{missing_field}'",
                    severity="warning"
                )
                result.warnings.append(warning)


class DataValidator:
    """数据验证器 - 验证 YAML 数据文件的完整性"""
    
    def __init__(self, schema_path=None):
        """初始化数据验证器
        
        Args:
            schema_path: Schema 文件路径（可选）。如果提供，将使用 schema 验证
        """
        self.errors = []
        self.warnings = []
        self.schema_validator = None
        
        # 如果提供了 schema_path，初始化 SchemaValidator
        if schema_path:
            try:
                self.schema_validator = SchemaValidator(schema_path)
            except Exception as e:
                # Schema 加载失败不影响基本验证功能
                print(f"警告: 无法加载 Schema: {str(e)}")
                self.schema_validator = None
    
    def validate_yaml_syntax(self, file_path):
        """验证 YAML 文件语法是否正确
        
        Returns:
            tuple: (is_valid, data_or_error)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
                if data is None:
                    return False, "YAML 文件为空"
                if not isinstance(data, dict):
                    return False, "YAML 根节点必须是字典类型"
                return True, data
        except yaml.YAMLError as e:
            return False, f"YAML 语法错误: {str(e)}"
        except FileNotFoundError:
            return False, "文件不存在"
        except Exception as e:
            return False, f"读取文件错误: {str(e)}"
    
    def extract_placeholders(self, doc):
        """从模板文档中提取所有占位符
        
        Args:
            doc: python-docx Document 对象
            
        Returns:
            set: 占位符键名集合（不含 {{}} 符号）
        """
        placeholders = set()
        pattern = r'\{\{([^}]+)\}\}'
        
        # 从段落中提取
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # 从表格中提取
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
        
        return placeholders
    
    def validate_required_keys(self, data, required_keys):
        """验证数据中是否包含所有必需的键
        
        Args:
            data: YAML 数据（字典）
            required_keys: 必需的键集合
            
        Returns:
            tuple: (is_valid, missing_keys, extra_keys)
        """
        data_keys = set(self._get_all_keys(data))
        missing_keys = required_keys - data_keys
        extra_keys = data_keys - required_keys
        
        is_valid = len(missing_keys) == 0
        return is_valid, missing_keys, extra_keys
    
    def _get_all_keys(self, data, prefix=''):
        """递归获取所有键（包括嵌套键，使用点号表示）"""
        keys = []
        if isinstance(data, dict):
            for key, value in data.items():
                full_key = f"{prefix}.{key}" if prefix else key
                keys.append(full_key)
                if isinstance(value, dict):
                    keys.extend(self._get_all_keys(value, full_key))
        return keys
    
    def validate_single_file(self, data_file, template_path):
        """验证单个数据文件
        
        Returns:
            dict: 验证结果
        """
        result = {
            "file": str(data_file),
            "valid": True,
            "errors": [],
            "warnings": []
        }
        
        # 1. 验证 YAML 语法
        is_valid, data_or_error = self.validate_yaml_syntax(data_file)
        if not is_valid:
            result["valid"] = False
            result["errors"].append({
                "type": "yaml_syntax",
                "message": data_or_error,
                "severity": "error"
            })
            return result
        
        data = data_or_error
        
        # 2. 如果有 SchemaValidator，使用 schema 验证（优先）
        if self.schema_validator:
            try:
                # 调用 SchemaValidator 进行验证
                schema_result = self.schema_validator.validate_against_schema(
                    data, 
                    file_path=str(data_file)
                )
                
                # 将 ValidationResult 转换为字典并合并结果
                result["errors"].extend([e.to_dict() for e in schema_result.errors])
                result["warnings"].extend([w.to_dict() for w in schema_result.warnings])
                
                if not schema_result.is_valid:
                    result["valid"] = False
                
                # 添加 schema 验证的额外信息
                result["schema_validation"] = {
                    "enabled": True,
                    "checked_fields": schema_result.checked_fields,
                    "schema_version": schema_result.schema_version,
                    "timestamp": schema_result.timestamp.isoformat()
                }
                
            except SchemaNotLoadedError as e:
                result["warnings"].append({
                    "type": "schema_not_loaded",
                    "message": f"Schema 未加载: {str(e)}",
                    "severity": "warning"
                })
            except SchemaValidationError as e:
                result["warnings"].append({
                    "type": "schema_validation_error",
                    "message": f"Schema 验证失败: {str(e)}",
                    "severity": "warning"
                })
            except Exception as e:
                result["warnings"].append({
                    "type": "unexpected_error",
                    "message": f"验证过程发生错误: {str(e)}",
                    "severity": "warning"
                })
        else:
            # 3. 回退到基于模板的验证（向后兼容）
            try:
                from docx import Document
                doc = Document(template_path)
                required_keys = self.extract_placeholders(doc)
            except Exception as e:
                result["valid"] = False
                result["errors"].append({
                    "type": "template_error",
                    "message": f"无法读取模板: {str(e)}",
                    "severity": "error"
                })
                return result
            
            # 验证必需键
            is_valid, missing_keys, extra_keys = self.validate_required_keys(data, required_keys)
            
            if missing_keys:
                result["valid"] = False
                for key in missing_keys:
                    result["errors"].append({
                        "type": "missing_key",
                        "message": f"缺少必需的键: '{key}'",
                        "key": key,
                        "severity": "error"
                    })
            
            if extra_keys:
                for key in extra_keys:
                    result["warnings"].append({
                        "type": "extra_key",
                        "message": f"未使用的数据键: '{key}'",
                        "key": key,
                        "severity": "warning"
                    })
            
            result["schema_validation"] = {
                "enabled": False
            }
        
        return result
    
    def validate_batch(self, data_dir, template_path):
        """批量验证数据文件
        
        Returns:
            dict: 批量验证结果
        """
        dir_path = Path(data_dir)
        yaml_files = sorted(list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml')))
        
        if not yaml_files:
            return {
                "success": False,
                "error": f"未在目录中找到 YAML 文件: {data_dir}",
                "total": 0,
                "valid": 0,
                "invalid": 0
            }
        
        results = []
        valid_count = 0
        invalid_count = 0
        
        for data_file in yaml_files:
            result = self.validate_single_file(data_file, template_path)
            results.append(result)
            
            if result["valid"]:
                valid_count += 1
            else:
                invalid_count += 1
        
        return {
            "success": invalid_count == 0,
            "total": len(yaml_files),
            "valid": valid_count,
            "invalid": invalid_count,
            "results": results
        }


class DocumentGenerator:
    def __init__(self):
        self.debug = False
    
    def log(self, message):
        """调试日志输出"""
        if self.debug:
            print(f"[DEBUG] {message}")
    
    def copy_paragraph_format(self, source_paragraph, target_paragraph):
        """安全地复制段落格式"""
        try:
            # 复制段落格式
            if source_paragraph.paragraph_format:
                target_format = target_paragraph.paragraph_format
                source_format = source_paragraph.paragraph_format
                
                # 复制对齐方式
                if source_format.alignment is not None:
                    target_format.alignment = source_format.alignment
                
                # 复制缩进
                if source_format.left_indent is not None:
                    target_format.left_indent = source_format.left_indent
                if source_format.right_indent is not None:
                    target_format.right_indent = source_format.right_indent
                if source_format.first_line_indent is not None:
                    target_format.first_line_indent = source_format.first_line_indent
                
                # 复制间距
                if source_format.space_before is not None:
                    target_format.space_before = source_format.space_before
                if source_format.space_after is not None:
                    target_format.space_after = source_format.space_after
                if source_format.line_spacing is not None:
                    target_format.line_spacing = source_format.line_spacing
            
            # 复制样式
            if source_paragraph.style:
                target_paragraph.style = source_paragraph.style
                
            self.log(f"成功复制段落格式")
            
        except Exception as e:
            self.log(f"复制段落格式时出错: {str(e)}")
    
    def get_base_run_format(self, paragraph):
        """获取段落的基础文字格式"""
        base_format = {}
        
        try:
            # 优先从现有run获取格式（更准确）
            if paragraph.runs:
                first_run = paragraph.runs[0]
                
                # 字体名称和大小
                if first_run.font.name:
                    base_format['font_name'] = first_run.font.name
                if first_run.font.size:
                    base_format['font_size'] = first_run.font.size
                
                # 获取颜色（更健壮的方式）
                if first_run.font.color:
                    if first_run.font.color.rgb:
                        base_format['font_color'] = first_run.font.color.rgb
                    elif first_run.font.color.theme_color:
                        # 如果是主题颜色，也尝试获取
                        base_format['font_color_theme'] = first_run.font.color.theme_color
                
                # 保留粗体、斜体、下划线状态（包括 None）
                base_format['font_bold'] = first_run.font.bold
                base_format['font_italic'] = first_run.font.italic
                base_format['font_underline'] = first_run.font.underline
                
                # 保留其他格式
                base_format['font_strike'] = first_run.font.strike
                base_format['font_all_caps'] = first_run.font.all_caps
                base_format['font_small_caps'] = first_run.font.small_caps
            
            # 如果从run获取不到，再从段落样式获取
            if not base_format and paragraph.style and hasattr(paragraph.style, 'font'):
                style_font = paragraph.style.font
                if style_font.name:
                    base_format['font_name'] = style_font.name
                if style_font.size:
                    base_format['font_size'] = style_font.size
                if style_font.color and style_font.color.rgb:
                    base_format['font_color'] = style_font.color.rgb
            
            self.log(f"获取到的基础格式: {base_format}")
        
        except Exception as e:
            self.log(f"获取基础格式时出错: {str(e)}")
        
        return base_format
    
    def apply_run_format(self, run, base_format):
        """应用基础格式到run"""
        try:
            # 字体名称和大小
            if 'font_name' in base_format:
                font_name = base_format['font_name']
                run.font.name = font_name
                
                # 🔑 关键修复：同时设置东亚字体（中文）和复杂脚本字体
                from docx.oxml.shared import OxmlElement
                from docx.oxml.ns import qn
                
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                
                # 设置所有字体类型
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:eastAsia'), font_name)  # 东亚字体（中文）
                rFonts.set(qn('w:cs'), font_name)        # 复杂脚本字体
            
            if 'font_size' in base_format:
                run.font.size = base_format['font_size']
            
            # 颜色
            if 'font_color' in base_format:
                run.font.color.rgb = base_format['font_color']
            
            # 粗体、斜体、下划线（保留 None 状态）
            if 'font_bold' in base_format:
                run.font.bold = base_format['font_bold']
            if 'font_italic' in base_format:
                run.font.italic = base_format['font_italic']
            if 'font_underline' in base_format:
                run.font.underline = base_format['font_underline']
            
            # 其他格式
            if 'font_strike' in base_format:
                run.font.strike = base_format['font_strike']
            if 'font_all_caps' in base_format:
                run.font.all_caps = base_format['font_all_caps']
            if 'font_small_caps' in base_format:
                run.font.small_caps = base_format['font_small_caps']
                
        except Exception as e:
            self.log(f"应用run格式时出错: {str(e)}")
    
    def parse_markdown_formatting(self, text):
        """解析Markdown格式并返回格式化的片段列表"""
        segments = []
        
        # 处理粗体和斜体的正则表达式
        # 支持 **粗体** 和 *斜体*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        
        parts = re.findall(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # 粗体文本
                segments.append({
                    'text': part[2:-2],  # 去掉 ** 标记
                    'bold': True,
                    'italic': False
                })
            elif part.startswith('*') and part.endswith('*'):
                # 斜体文本
                segments.append({
                    'text': part[1:-1],  # 去掉 * 标记
                    'bold': False,
                    'italic': True
                })
            else:
                # 普通文本
                segments.append({
                    'text': part,
                    'bold': False,
                    'italic': False
                })
        
        return segments
    
    def add_formatted_text_to_paragraph(self, paragraph, text, base_format):
        """向段落添加格式化文本"""
        # 清空现有内容
        paragraph.clear()
        
        # 按行分割文本（保持 \n 作为换行）
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if i > 0:
                # 添加换行符，并应用基础格式（保持缩进和字体）
                newline_run = paragraph.add_run('\n')
                self.apply_run_format(newline_run, base_format)
            
            # 解析Markdown格式
            segments = self.parse_markdown_formatting(line)
            
            for segment in segments:
                run = paragraph.add_run(segment['text'])
                
                # 应用基础格式
                self.apply_run_format(run, base_format)
                
                # 应用Markdown格式（只在有 Markdown 标记时覆盖）
                if segment['bold']:
                    run.font.bold = True
                if segment['italic']:
                    run.font.italic = True
    
    def process_template_paragraph(self, paragraph, data):
        """处理单个模板段落"""
        if not paragraph.text.strip():
            return False
        
        original_text = paragraph.text
        self.log(f"处理段落: {original_text[:50]}...")
        
        # ⚠️ 关键：在 clear() 之前获取格式，因为 clear() 会删除所有 runs
        base_format = self.get_base_run_format(paragraph)
        
        # 替换占位符
        processed_text = self.replace_placeholders(original_text, data)
        
        if processed_text != original_text:
            self.log(f"文本已替换: {processed_text[:50]}...")
            
            # 检查是否需要分割段落（\n\n 表示段落分割）
            if '\n\n' in processed_text:
                return self.split_into_multiple_paragraphs(paragraph, processed_text, base_format, data)
            else:
                # 单段落处理
                self.add_formatted_text_to_paragraph(paragraph, processed_text, base_format)
                return True
        
        return False
    
    def split_into_multiple_paragraphs(self, original_paragraph, text, base_format, data):
        """将文本分割成多个段落"""
        # 按 \n\n 分割段落
        paragraphs_text = text.split('\n\n')
        
        # 处理第一个段落（替换原段落）
        if paragraphs_text:
            first_text = paragraphs_text[0].strip()
            if first_text:
                self.add_formatted_text_to_paragraph(original_paragraph, first_text, base_format)
        
        # 检查段落是否在表格单元格中
        parent = original_paragraph._element.getparent()
        
        # 如果段落在表格单元格中（parent 是 tc 元素）
        if parent.tag.endswith('}tc'):
            # 在表格单元格中，直接在当前单元格添加新段落
            from docx.oxml import OxmlElement
            from docx.text.paragraph import Paragraph
            
            for para_text in paragraphs_text[1:]:
                para_text = para_text.strip()
                if para_text:
                    # 在单元格中添加新段落
                    new_p_element = OxmlElement('w:p')
                    parent.append(new_p_element)
                    
                    # 创建 Paragraph 对象
                    new_paragraph = Paragraph(new_p_element, parent)
                    
                    # 复制格式
                    self.copy_paragraph_format(original_paragraph, new_paragraph)
                    
                    # 添加格式化文本
                    self.add_formatted_text_to_paragraph(new_paragraph, para_text, base_format)
        else:
            # 在普通文档流中，按原有逻辑处理
            original_index = list(parent).index(original_paragraph._element)
            
            for i, para_text in enumerate(paragraphs_text[1:], 1):
                para_text = para_text.strip()
                if para_text:
                    # 获取文档对象
                    if hasattr(original_paragraph, '_part'):
                        doc_obj = original_paragraph._part.document
                    else:
                        # 如果无法获取文档对象，跳过
                        self.log("无法获取文档对象，跳过段落创建")
                        continue
                    
                    # 创建新段落
                    new_paragraph = doc_obj.add_paragraph()
                    
                    # 复制格式
                    self.copy_paragraph_format(original_paragraph, new_paragraph)
                    
                    # 添加格式化文本
                    self.add_formatted_text_to_paragraph(new_paragraph, para_text, base_format)
                    
                    # 将新段落插入到正确位置
                    parent.insert(original_index + i, new_paragraph._element)
        
        return True
    
    def replace_placeholders(self, text, data):
        """替换文本中的占位符"""
        def replace_match(match):
            placeholder = match.group(1)
            return self.get_nested_value(data, placeholder)
        
        # 匹配 {{key}} 格式的占位符
        result = re.sub(r'\{\{([^}]+)\}\}', replace_match, text)
        return result
    
    def get_nested_value(self, data, key_path):
        """获取嵌套字典中的值"""
        keys = key_path.strip().split('.')
        value = data
        
        try:
            for key in keys:
                if isinstance(value, dict):
                    value = value.get(key, f'{{{{{key_path}}}}}')
                else:
                    return f'{{{{{key_path}}}}}'
            
            # 如果值是列表，转换为换行分隔的字符串
            if isinstance(value, list):
                return '\n'.join(str(item) for item in value)
            
            return str(value) if value is not None else f'{{{{{key_path}}}}}'
        
        except Exception as e:
            self.log(f"获取值时出错 {key_path}: {str(e)}")
            return f'{{{{{key_path}}}}}'
    
    def analyze_directory(self, directory_path, recursive=False):
        """分析目录中的YAML文件数量"""
        try:
            dir_path = Path(directory_path)
            
            if not dir_path.exists():
                return {
                    "success": False,
                    "error": f"目录不存在: {directory_path}"
                }
            
            if not dir_path.is_dir():
                return {
                    "success": False,
                    "error": f"路径不是目录: {directory_path}"
                }
            
            # 查找YAML文件
            if recursive:
                yaml_files = list(dir_path.rglob('*.yml')) + list(dir_path.rglob('*.yaml'))
            else:
                yaml_files = list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml'))
            
            return {
                "success": True,
                "file_count": len(yaml_files),
                "files": [str(f.relative_to(dir_path)) for f in yaml_files]
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def batch_generate(self, data_dir, template_path, output_dir, continue_on_error=False, debug=False, quiet=False, verbose=False):
        """批量生成文档
        
        Args:
            data_dir: 数据文件目录
            template_path: 模板文件路径
            output_dir: 输出目录
            continue_on_error: 遇错继续
            debug: 调试模式
            quiet: 安静模式（仅显示错误）
            verbose: 详细模式（显示详细日志）
        """
        self.debug = debug
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 获取所有YAML文件
        dir_path = Path(data_dir)
        yaml_files = sorted(list(dir_path.glob('*.yml')) + list(dir_path.glob('*.yaml')))
        
        if not yaml_files:
            return {
                "success": False,
                "error": f"未在目录中找到YAML文件: {data_dir}",
                "total": 0,
                "succeeded": 0,
                "failed": 0
            }
        
        # 批量处理
        results = []
        succeeded = 0
        failed = 0
        
        # 进度条设置
        use_progress_bar = not quiet and not verbose
        
        if not quiet:
            print(f"开始批量生成文档...")
            print(f"总共发现 {len(yaml_files)} 个数据文件")
            if not use_progress_bar:
                print("-" * 60)
        
        # 使用 tqdm 进度条或普通循环
        if use_progress_bar:
            pbar = tqdm(yaml_files, desc="批量生成文档", unit="file")
            iterator = pbar
        else:
            iterator = yaml_files
            pbar = None
        
        for idx, data_file in enumerate(iterator, 1):
            file_name = data_file.stem
            output_path = os.path.join(output_dir, f"{file_name}.docx")
            
            if verbose and not use_progress_bar:
                print(f"\n[{idx}/{len(yaml_files)}] 处理: {data_file.name}")
            elif not use_progress_bar and not quiet:
                print(f"\n[{idx}/{len(yaml_files)}] 处理: {data_file.name}")
            
            try:
                success = self.generate_document(
                    template_path,
                    str(data_file),
                    output_path,
                    debug=False  # 批量模式下关闭详细日志
                )
                
                if success:
                    succeeded += 1
                    results.append({
                        "file": data_file.name,
                        "status": "success",
                        "output": output_path
                    })
                    # Only show success message in verbose mode without progress bar
                    if verbose and not use_progress_bar:
                        print(f"  ✅ 成功: {output_path}")
                else:
                    failed += 1
                    results.append({
                        "file": data_file.name,
                        "status": "failed",
                        "error": "生成失败"
                    })
                    if not quiet:
                        print(f"  ❌ 失败: {data_file.name}")
                    if not continue_on_error:
                        break
                        
            except Exception as e:
                failed += 1
                error_msg = str(e)
                if not quiet:
                    print(f"  ❌ 错误: {error_msg}")
                results.append({
                    "file": data_file.name,
                    "status": "failed",
                    "error": error_msg
                })
                
                if not continue_on_error:
                    break
        
        # 生成汇总报告
        if not quiet:
            print("\n" + "=" * 60)
            print("批量生成完成!")
            print(f"  ✅ 成功: {succeeded}")
            print(f"  ❌ 失败: {failed}")
            print(f"  📊 总计: {len(yaml_files)}")
            print("=" * 60)
        
        return {
            "success": failed == 0,
            "total": len(yaml_files),
            "succeeded": succeeded,
            "failed": failed,
            "results": results
        }
    
    def generate_document(self, template_path, data_path, output_path, debug=False):
        """生成文档的主函数"""
        self.debug = debug
        
        try:
            # 加载数据
            self.log(f"加载数据文件: {data_path}")
            with open(data_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
            
            # 加载模板
            self.log(f"加载模板文件: {template_path}")
            doc = Document(template_path)
            
            # 处理所有段落
            processed_count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                if self.process_template_paragraph(paragraph, data):
                    processed_count += 1
            
            self.log(f"共处理了 {processed_count} 个段落")
            
            # 处理所有表格
            table_cell_count = 0
            for table_idx, table in enumerate(doc.tables):
                self.log(f"处理表格 {table_idx}")
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        # 处理单元格中的每个段落
                        for paragraph in cell.paragraphs:
                            if self.process_template_paragraph(paragraph, data):
                                table_cell_count += 1
            
            self.log(f"共处理了 {table_cell_count} 个表格单元格")
            
            # 保存文档
            self.log(f"保存文档到: {output_path}")
            doc.save(output_path)
            
            if debug:
                print(f"文档已成功生成: {output_path}")
                print(f"  - 处理段落: {processed_count}")
                print(f"  - 处理表格单元格: {table_cell_count}")
            
            return True
            
        except Exception as e:
            if debug:
                print(f"生成文档时出错: {str(e)}")
            return False

def cmd_validate(args):
    """validate 命令处理函数"""
    # 确定 schema 路径
    schema_path = getattr(args, 'schema', None)
    
    # 如果没有指定 schema，尝试使用默认路径
    if not schema_path:
        default_schema = Path(__file__).parent / "schemas" / "lesson_data_schema.yml"
        if default_schema.exists():
            schema_path = str(default_schema)
    
    # 初始化验证器（带或不带 schema）
    validator = DataValidator(schema_path=schema_path)
    
    # 输出验证模式信息
    if schema_path and validator.schema_validator:
        print(f"ℹ️  使用 Schema 验证: {schema_path}")
    else:
        print("ℹ️  使用模板占位符验证（未启用 Schema 验证）")
    
    if args.data_file:
        # 验证单个文件
        result = validator.validate_single_file(args.data_file, args.template)
        
        # 输出结果
        if result["valid"]:
            print(f"✅ 验证通过: {args.data_file}")
            if result["warnings"]:
                print(f"\n⚠️  警告 ({len(result['warnings'])}):")
                for warning in result["warnings"]:
                    print(f"  - {warning['message']}")
        else:
            print(f"❌ 验证失败: {args.data_file}")
            print(f"\n错误 ({len(result['errors'])}):")
            for error in result["errors"]:
                print(f"  - {error['message']}")
            
            if result["warnings"]:
                print(f"\n⚠️  警告 ({len(result['warnings'])}):")
                for warning in result["warnings"]:
                    print(f"  - {warning['message']}")
        
        # JSON 输出
        print("\n" + json.dumps(result, ensure_ascii=False, indent=2))
        
        return 0 if result["valid"] else 1
    
    else:
        # 批量验证
        result = validator.validate_batch(args.data_dir, args.template)
        
        if not result.get("success", False) and "error" in result:
            print(f"❌ 错误: {result['error']}")
            print(json.dumps(result, ensure_ascii=False, indent=2))
            return 1
        
        # 输出汇总
        print("=" * 60)
        print("数据验证完成!")
        print(f"  ✅ 有效: {result['valid']}")
        print(f"  ❌ 无效: {result['invalid']}")
        print(f"  📊 总计: {result['total']}")
        print("=" * 60)
        
        # 显示详细错误
        if result['invalid'] > 0:
            print("\n详细错误:")
            for file_result in result['results']:
                if not file_result['valid']:
                    print(f"\n❌ {file_result['file']}:")
                    for error in file_result['errors']:
                        print(f"   - {error['message']}")
        
        # JSON 输出
        print("\n" + "=" * 60)
        print("详细结果 (JSON):")
        print(json.dumps(result, ensure_ascii=False, indent=2))
        
        return 0 if result['success'] else 1

def cmd_analyze(args):
    """analyze 命令处理函数"""
    generator = DocumentGenerator()
    result = generator.analyze_directory(args.directory, recursive=args.recursive)
    
    # 输出JSON格式结果
    print(json.dumps(result, ensure_ascii=False, indent=2))
    
    return 0 if result['success'] else 1

def cmd_generate(args):
    """generate 命令处理函数"""
    generator = DocumentGenerator()
    
    # 确定输出路径
    if args.output:
        output_path = args.output
    else:
        template_path = Path(args.template)
        data_path = Path(args.data)
        output_dir = args.output_dir if args.output_dir else './output'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"{data_path.stem}.docx")
    
    # 生成文档
    success = generator.generate_document(
        args.template, 
        args.data, 
        output_path, 
        debug=args.debug
    )
    
    # 输出结果（非debug模式）
    if success and not args.debug:
        print(f"文档已成功生成: {output_path}")
    elif not success:
        print(f"文档生成失败")
    
    # 输出JSON格式结果
    result = {
        "success": success,
        "output_path": output_path if success else None
    }
    print(json.dumps(result, ensure_ascii=False))
    
    return 0 if success else 1

def cmd_batch(args):
    """batch 命令处理函数"""
    generator = DocumentGenerator()
    
    # 批量生成
    result = generator.batch_generate(
        args.data_dir,
        args.template,
        args.output_dir,
        continue_on_error=args.continue_on_error,
        debug=args.debug,
        quiet=args.quiet,
        verbose=args.verbose
    )
    
    # 输出JSON格式的详细结果
    if not args.quiet:
        print("\n" + "=" * 60)
        print("详细结果 (JSON):")
        print(json.dumps(result, ensure_ascii=False, indent=2))
    
    return 0 if result['success'] else 1

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='Docu-Weaver: 自动化文档生成工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 分析目录中的YAML文件
  python generate_docs.py analyze ./test_data
  
  # 验证单个数据文件
  python generate_docs.py validate lesson1.yml template.docx
  
  # 批量验证数据文件
  python generate_docs.py validate --batch ./test_data template.docx
  
  # 生成单个文档
  python generate_docs.py generate lesson1.yml template.docx output/
  
  # 批量生成文档
  python generate_docs.py batch ./test_data template.docx ./output
  
  # 批量生成（遇错继续）
  python generate_docs.py batch ./data template.docx ./output --continue-on-error
  
  # 递归分析子目录
  python generate_docs.py analyze ./data --recursive
        """
    )
    
    # 创建子命令
    subparsers = parser.add_subparsers(dest='command', help='可用命令')
    
    # analyze 子命令
    parser_analyze = subparsers.add_parser('analyze', help='分析目录中的YAML文件数量')
    parser_analyze.add_argument('directory', help='要分析的目录路径')
    parser_analyze.add_argument('-r', '--recursive', action='store_true', 
                               help='递归扫描子目录')
    parser_analyze.set_defaults(func=cmd_analyze)
    
    # validate 子命令
    parser_validate = subparsers.add_parser('validate', help='验证数据文件完整性')
    parser_validate.add_argument('data_file', nargs='?', help='YAML数据文件路径（单文件模式）')
    parser_validate.add_argument('template', help='Word模板文件路径')
    parser_validate.add_argument('--batch', dest='data_dir', help='数据文件目录路径（批量模式）')
    parser_validate.add_argument('--schema', help='Schema 文件路径（可选，默认使用 schemas/lesson_data_schema.yml）')
    parser_validate.set_defaults(func=cmd_validate)
    
    # generate 子命令
    parser_generate = subparsers.add_parser('generate', help='生成单个文档')
    parser_generate.add_argument('data', help='YAML数据文件路径')
    parser_generate.add_argument('template', help='Word模板文件路径')
    parser_generate.add_argument('output_dir', help='输出目录路径')
    parser_generate.add_argument('-o', '--output', help='指定输出文件完整路径（可选）')
    parser_generate.add_argument('--debug', action='store_true', help='启用调试模式')
    parser_generate.set_defaults(func=cmd_generate)
    
    # batch 子命令
    parser_batch = subparsers.add_parser('batch', help='批量生成文档')
    parser_batch.add_argument('data_dir', help='数据文件目录路径')
    parser_batch.add_argument('template', help='Word模板文件路径')
    parser_batch.add_argument('output_dir', help='输出目录路径')
    parser_batch.add_argument('--continue-on-error', action='store_true',
                             help='遇到错误时继续处理其他文件')
    parser_batch.add_argument('--debug', action='store_true', help='启用调试模式')
    parser_batch.add_argument('-q', '--quiet', action='store_true', 
                             help='安静模式（仅显示错误）')
    parser_batch.add_argument('-v', '--verbose', action='store_true', 
                             help='详细模式（显示所有处理信息）')
    parser_batch.set_defaults(func=cmd_batch)
    
    # 解析参数
    args = parser.parse_args()
    
    # 如果没有指定子命令，显示帮助信息
    if not args.command:
        parser.print_help()
        return 1
    
    # 执行对应的命令函数
    return args.func(args)

if __name__ == "__main__":
    exit(main())
